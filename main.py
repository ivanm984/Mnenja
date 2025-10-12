# main.py
from dotenv import load_dotenv
load_dotenv()

import os
import json
import re
import io
from typing import List, Optional, Dict, Any, Tuple
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse
from pydantic import BaseModel, Field
from datetime import datetime

# AI - Gemini
import google.generativeai as genai

# PDF parsing
from pypdf import PdfReader
import fitz  # PyMuPDF
from PIL import Image

# DOCX
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches

# =============================================================================
# KONFIGURACIJA
# =============================================================================

API_KEY = os.environ.get("GEMINI_API_KEY")
if not API_KEY:
    raise RuntimeError("❌ GEMINI_API_KEY manjka v .env datoteki!")

print(f"✅ API Key najden: {API_KEY[:10]}...")

genai.configure(api_key=API_KEY)

MODEL_NAME = "gemini-2.5-pro"
GEN_CFG = {
    "temperature": 0.0,
    "top_p": 0.9,
    "top_k": 40,
    "max_output_tokens": 40000,
    "response_mime_type": "application/json",
}

# =============================================================================
# NALAGANJE BAZE ZNANJA (OPN, PRILOGE IN IZRAZI)
# =============================================================================

KEYWORD_TO_CLEN = {
    # Gradnja in objekti
    "gradnj": "52_clen", "dozidava": "52_clen", "nadzidava": "52_clen", "rekonstrukcija": "52_clen",
    "odstranitev": "52_clen", "sprememba namembnosti": "54_clen", "vrste objektov": "56_clen",
    "nezahtevni objekt": "64_clen", "enostavni objekt": "64_clen", "razpršena gradnja": "102_clen",
    "nelegalna gradnja": "103_clen",

    # Urejanje parcele
    "odmik": "58_clen", "odmiki": "58_clen", "soglasje soseda": "58_clen", "regulacijsk": "57_clen",
    "velikost parcele": "66_clen", "parcela objekta": "66_clen",
    "velikost objektov": "59_clen", "faktor izrabe": "59_clen", "FI": "59_clen", "faktor zazidanosti": "59_clen", "FZ": "59_clen",
    "višina objekt": "59_clen",

    # Oblikovanje
    "oblikovanj": "60_clen", "fasad": "60_clen", "streh": "60_clen", "kritina": "60_clen", "naklon strehe": "60_clen",
    "zelene površine": "61_clen", "FZP": "61_clen", "igrišče": "61_clen",

    # Infrastruktura
    "parkirišč": "62_clen", "parkirna mesta": "62_clen", "garaž": "62_clen", "število parkirnih mest": "63_clen",
    "komunaln": "67_clen", "priključek": "69_clen", "priključitev": "69_clen",
    "vodovod": "73_clen", "kanalizacij": "74_clen", "greznica": "69_clen", "čistilna naprava": "74_clen",
    "plinovod": "76_clen", "elektro": "77_clen", "daljnovod": "77_clen", "javna razsvetljava": "78_clen",
    "telekomunikacijsk": "79_clen", "komunikacijsk": "79_clen",

    # Varovanje in omejitve
    "varovalni pas": "70_clen", "varstvo narave": "81_clen", "kulturna dediščina": "82_clen",
    "vplivi na okolje": "83_clen", "varstvo voda": "85_clen", "vodotok": "85_clen",
    "priobalnem zemljišču": "85_clen", "vodovarstven": "86_clen",
    "varovalni gozd": "88_clen", "gozd s posebnim namenom": "89_clen",
    "hrup": "98_clen", "sevanje": "99_clen", "osončenj": "100_clen",
    "poplavn": "94_clen", "erozij": "92_clen", "plaz": "92_clen", "plazljiv": "92_clen",
    "potresn": "93_clen", "požar": "95_clen",
    
    # Ostalo
    "oglaševanj": "65_clen", "odpadk": "80_clen", "mineralne surovine": "90_clen",
    "obrambne potrebe": "96_clen", "zaklonišč": "96_clen",
    "invalid": "97_clen", "dostop za invalide": "97_clen", "arhitektonske ovire": "97_clen"
}


def format_structured_content(data_dict: Dict[str, Any]) -> str:
    """Pretvori strukturiran JSON objekt v berljivo besedilo za AI prompt."""
    lines = []
    for key, value in data_dict.items():
        if isinstance(value, dict):
            lines.append(f"\n- {key.replace('_', ' ').capitalize()}:")
            for sub_key, sub_value in value.items():
                lines.append(f"  - {sub_key.replace('_', ' ')}: {sub_value}")
        elif isinstance(value, list):
            lines.append(f"\n- {key.replace('_', ' ').capitalize()}:")
            for item in value:
                lines.append(f"  - {item}")
        else:
            lines.append(f"- {key.replace('_', ' ').capitalize()}: {value}")
    return "\n".join(lines)

def format_uredba_summary(uredba_data: Dict[str, Any]) -> str:
    """Pretvori podatke iz UredbaObjekti.json v tekstovno predstavitev."""
    if not uredba_data:
        return "Podatki iz UredbaObjekti.json niso na voljo."
    try:
        return json.dumps(uredba_data, ensure_ascii=False, indent=2)
    except Exception:
        return str(uredba_data)

def load_knowledge_base() -> Tuple[Dict, Dict, List, Dict, str, str]:
    """Naloži celotno bazo znanja: OPN, priloge in slovar izrazov."""
    try:
        with open("OPN.json", "r", encoding="utf-8") as f: opn_katalog = json.load(f)

        clen_data_map = {}
        for cat_key, cat_data in opn_katalog.items():
            if 'clen' in cat_data and 'podrocja' in cat_data and isinstance(cat_data['podrocja'], dict):
                for raba_key, raba_data in cat_data['podrocja'].items():
                    clen_data_map[raba_key.upper()] = {"title": cat_data.get("naslov", ""),"podrocje_naziv": raba_data.get("naziv", ""),"content_structured": raba_data,"parent_clen_key": f"{cat_data['clen']}_clen"}
        print(f"✅ Ustvarjeno kazalo za {len(clen_data_map)} podrobnih namenskih rab.")
        
        priloge = {}
        for f in ["priloga1.json", "priloga2.json", "priloga3-4.json", "Izrazi.json"]:
            try:
                with open(f, "r", encoding="utf-8") as file:
                    if f == "priloga3-4.json":
                        data_3_4 = json.load(file)
                        priloge.update({'priloga3': data_3_4.get('priloga3', {}), 'priloga4': data_3_4.get('priloga4', {})})
                    else:
                        priloge[f.split('.')[0]] = json.load(file)
                print(f"✅ Uspešno naložena datoteka: {f}")
            except (FileNotFoundError, json.JSONDecodeError) as e:
                print(f"⚠️ Napaka pri nalaganju {f}: {e}.")
                if f == "priloga3-4.json": priloge.update({'priloga3': {}, 'priloga4': {}})
                else: priloge[f.split('.')[0]] = {}

        try:
            with open("UredbaObjekti.json", "r", encoding="utf-8") as f:
                uredba_data = json.load(f)
            print("✅ Uspešno naložena datoteka: UredbaObjekti.json")
        except (FileNotFoundError, json.JSONDecodeError) as e:
            print(f"⚠️ Napaka pri nalaganju UredbaObjekti.json: {e}.")
            uredba_data = {}

        all_eups = [item.get("enota_urejanja", "") for item in priloge.get('priloga2', {}).get('table_entries', [])]
        all_eups.extend([item.get("urejevalna_enota", "") for item in priloge.get('priloga3', {}).get('entries', [])])
        unique_eups = sorted(list(set(filter(None, all_eups))), key=len, reverse=True)

        izrazi_data = priloge.get('Izrazi', {})
        izrazi_text = "\n".join([f"- **{term['term']}**: {term['definition']}" for term in izrazi_data.get("terms", [])])

        uredba_text = format_uredba_summary(uredba_data)

        print("✅ Baza znanja uspešno naložena.")
        return opn_katalog, priloge, unique_eups, clen_data_map, izrazi_text, uredba_text
    except Exception as e:
        raise RuntimeError(f"❌ Kritična napaka pri nalaganju baze znanja: {e}")

OPN_KATALOG, PRILOGE, ALL_EUPS, CLEN_DATA_MAP, IZRAZI_TEXT, UREDBA_TEXT = load_knowledge_base()

# =============================================================================
# INTELIGENTNA ANALIZA IN SESTAVA ZAHTEV
# =============================================================================

def call_gemini_for_details(project_text: str, images: List[Image.Image]) -> Dict[str, Optional[List[str]]]:
    """1. AI Klic (Detektiv): Prvotno iz besedila in po potrebi še slik projekta izlušči EUP in namensko rabo."""
    print("🤖 1. Klic AI (Detektiv): Iščem EUP in namensko rabo (analiza besedila in po potrebi slik)...")
    prompt = f"""
    Analiziraj spodnje besedilo iz projektne dokumentacije. Če tega ne najdeš poglej slike. Tvoja naloga je najti dve informaciji:
    1.  **Enota Urejanja Prostora (EUP)**: To so oznake lokacij. Poišči VSE ustrezne oznake, ker lahko projekt (npr. vodovod) poteka preko več EUP. Bodi ekstremno natančen. Ne vračaj splošnih oznak, če so na voljo bolj specifične.
    2.  **Podrobnejša namenska raba**: To so kratice, npr. 'SSe', 'SK' ali 'A'. Poišči VSE, saj lahko projekt poteka preko več namenskih rab.

    Odgovori SAMO v JSON formatu, pri čemer sta vrednosti seznama (array). Primeri:
    {{ "eup": ["KE6-A6 549*", "KE6-A7 550*"], "namenska_raba": ["A", "SSe"] }}
    
    Če katerega od podatkov ne najdeš, vrni prazen seznam (array) za to polje.

    Besedilo dokumentacije: --- {project_text[:40000]} ---
    """
    try:
        model = genai.GenerativeModel(MODEL_NAME)
        
        content_parts = [prompt]
        if images:
            content_parts.extend(images)
            
        response = model.generate_content(content_parts)
        
        clean_response = re.sub(r'```(json)?', '', response.text, flags=re.IGNORECASE).strip()
        details = json.loads(clean_response)
        
        # Filtriranje in pretvorba: zagotovi, da sta EUP in namenska_raba seznama nizov
        eup_list = [str(e) for e in details.get('eup', []) if e]
        raba_list = [str(r).upper() for r in details.get('namenska_raba', []) if r]
        
        print(f"✅ AI Detektiv je našel: EUP={eup_list}, Namenska raba={raba_list}")
        return {"eup": eup_list, "namenska_raba": raba_list}
    except Exception as e:
        print(f"⚠️ Napaka pri AI Detektivu: {e}.")
        return {"eup": [], "namenska_raba": []}

def call_gemini_for_metadata(project_text: str) -> Dict[str, str]:
    """Dodaten AI klic za pridobivanje metapodatkov projekta."""
    print("🤖 1b. Klic AI (Arhivar): Iščem metapodatke projekta...")
    prompt = f"""
    Analiziraj priloženo besedilo projektne dokumentacije in izlušči naslednje podatke:
    1.  **ime_projekta**: Polno ime ali naziv projekta (npr. "Novogradnja enostanovanjske stavbe").
    2.  **stevilka_projekta**: Identifikacijska številka projekta.
    3.  **datum_projekta**: Datum izdelave dokumentacije.
    4.  **projektant**: Ime podjetja ali odgovornega projektanta.

    Odgovori SAMO v JSON formatu. Če katerega od podatkov ne najdeš, za vrednost uporabi "Ni podatka".
    Primer odgovora:
    {{
        "ime_projekta": "Prizidek in rekonstrukcija objekta",
        "stevilka_projekta": "P-123/2024",
        "datum_projekta": "Maj 2024",
        "projektant": "Projekt d.o.o."
    }}

    Besedilo dokumentacije:
    ---
    {project_text[:20000]}
    ---
    """
    try:
        model = genai.GenerativeModel(MODEL_NAME, generation_config={"response_mime_type": "application/json"})
        response = model.generate_content(prompt)
        clean_response = re.sub(r'```(json)?', '', response.text, flags=re.IGNORECASE).strip()
        metadata = json.loads(clean_response)
        print(f"✅ AI Arhivar je našel: {metadata}")
        return metadata
    except Exception as e:
        print(f"⚠️ Napaka pri AI Arhivistu: {e}.")
        return {
            "ime_projekta": "Ni podatka",
            "stevilka_projekta": "Ni podatka",
            "datum_projekta": "Ni podatka",
            "projektant": "Ni podatka"
        }

def call_gemini_for_key_data(project_text: str, images: List[Image.Image]) -> Dict[str, Any]:
    """POSODOBLJENO: Ciljana ekstrakcija ključnih tehničnih podatkov (iz besedila in slik)."""
    print("🤖 1c. Klic AI (Ekstraktor): Ciljana ekstrakcija ključnih tehničnih podatkov iz besedila IN grafike...")
    
    # Definiranje vseh 16 ključnih podatkovnih točk za ekstrakcijo
    KEY_DATA_PROMPT_MAP = {
        "glavni_objekt": "Natančen opis glavnega objekta (npr. enostanovanjska hiša, gospodarski objekt, opiši funkcijo).",
        "vrsta_gradnje": "Vrsta gradnje (npr. novogradnja, dozidava, nadzidava, rekonstrukcija, sprememba namembnosti).",
        "klasifikacija_cc_si": "CC-SI oziroma druga uradna klasifikacija objekta, če je navedena.",
        "nezahtevni_objekti": "Ali projekt vključuje nezahtevne objekte (navedi katere in njihove dimenzije).",
        "enostavni_objekti": "Ali projekt vključuje enostavne objekte (navedi katere in njihove dimenzije).",
        "vzdrzevalna_dela": "Opiši načrtovana vzdrževalna dela ali manjše rekonstrukcije, če so predvidene.",
        "parcela_objekta": "Številka gradbene/osnovne parcele (npr. 123/5).",
        "stevilke_parcel_ko": "Vse parcele in katastrska občina, ki so del projekta (npr. 123/5, 124/6, k.o. Litija).",
        "velikost_parcel": "Skupna velikost vseh parcel (npr. 1500 m2).",
        "velikost_obstojecega_objekta": "Velikost in etažnost obstoječih objektov na parceli (npr. hiša 10x8m P+1N, pomožni objekt 5x4m).",
        "tlorisne_dimenzije": "Zunanje tlorisne dimenzije NOVEGA glavnega objekta (npr. 12.0 m x 8.5 m).",
        "gabariti_etaznost": "Navedi etažnost in vertikalni gabarit NOVEGA objekta (npr. K+P+1N+M, višina kolenčnega zidu 1.5 m).",
        "faktor_zazidanosti_fz": "Vrednost faktorja zazidanosti (npr. 0.35 ali FZ=35%).",
        "faktor_izrabe_fi": "Vrednost faktorja izrabe (npr. 0.70 ali FI=0.7).",
        "zelene_povrsine": "Velikost in/ali faktor zelenih površin (npr. 700 m2, FZP=0.47).",
        "naklon_strehe": "Naklon strehe v stopinjah in tip (npr. 40° ali simetrična dvokapnica, 40 stopinj).",
        "kritina_barva": "Material in barva strešne kritine (npr. opečna kritina, temno rdeča).",
        "materiali_gradnje": "Tipični materiali (npr. masivna lesena hiša ali opeka, klasična gradnja).",
        "smer_slemena": "Orientacija slemena glede na plastice (npr. vzporedno s cesto/vrstnim redom gradnje).",
        "visinske_kote": "Pomembne kote (k.n.t., k.p. pritličja, k. slemena) (npr. k.p. = 345.50 m n.m.).",
        "odmiki_parcel": "Najmanjši in najpomembnejši navedeni odmiki od sosednjih parcelnih meja (npr. Južna meja: 4.5 m; Severna meja: 8.0 m).",
        "komunalni_prikljucki": "Opis priključitve na javno komunalno omrežje (elektrika, vodovod, kanalizacija itd.)."
    }

    prompt_items = "\n".join([f"{i+1}. **{key}**: {desc}" for i, (key, desc) in enumerate(KEY_DATA_PROMPT_MAP.items())])
    
    prompt = f"""
    Iz priložene projektne dokumentacije (besedila in slik - grafičnega dela) natančno izlušči naslednje ključne podatke projekta. Iščite numerične vrednosti in dimenzije!

    **Posebej bodite pozorni na informacije, ki se običajno nahajajo samo v grafikah/situaciji (odmiki, kote, dimenzije).**
    
    Odgovori SAMO v JSON formatu, pri čemer so vsi podatki nizi (string). Če podatka ni mogoče najti (ne v besedilu ne na slikah), uporabi vrednost "Ni podatka v dokumentaciji".

    ZAHTEVANI PODATKI:
    {prompt_items}

    Besedilo dokumentacije: --- {project_text[:40000]} ---
    """
    try:
        model = genai.GenerativeModel(MODEL_NAME)
        
        content_parts = [prompt]
        if images:
            content_parts.extend(images)
            
        response = model.generate_content(content_parts)
        
        clean_response = re.sub(r'```(json)?', '', response.text, flags=re.IGNORECASE).strip()
        key_data = json.loads(clean_response)
        
        print(f"✅ AI Ekstraktor je našel ključne podatke: {key_data}")
        
        # Zagotavljanje, da ima vračanje VSE ključe, ne glede na to, ali jih je AI našel
        final_data = {}
        for key in KEY_DATA_PROMPT_MAP.keys():
            final_data[key] = key_data.get(key, "Ni podatka v dokumentaciji")
            
        return final_data
        
    except Exception as e:
        print(f"⚠️ Napaka pri AI Ekstraktorju: {e}. Vračam prazne podatke.")
        # Vračanje vseh ključev z napako
        error_data = {}
        for key in KEY_DATA_PROMPT_MAP.keys():
            error_data[key] = "Napaka pri ekstrakciji"
        return error_data


def normalize_eup(eup_str: str) -> str:
    """Uporabljeno za popolno ujemanje."""
    if not eup_str: return ""
    return eup_str.strip().upper() 

def extract_referenced_namenske_rabe(content: str) -> List[str]:
    referenced = [m.upper() for p in [
        r'pogoj[ie]?\s+za\s+([A-Z]{1,3}[a-z]?)\b', r'kot\s+pri\s+([A-Z]{1,3}[a-z]?)\b',
        r'velj[a]?[jo]?\s+določila\s+za\s+([A-Z]{1,3}[a-z]?)\b', r'smiselno\s+velj[a]?[jo]?\s+za\s+([A-Z]{1,3}[a-z]?)\b',
        r'upoštev[a]?[jo]?\s+se\s+pogoj[ie]?\s+za\s+([A-Z]{1,3}[a-z]?)\b', r'skladno\s+s\s+pogoj[ie]?\s+za\s+([A-Z]{1,3}[a-z]?)\b',
        r'prevzem[a]?[jo]?\s+določila\s+za\s+([A-Z]{1,3}[a-z]?)\b', r'določila\s+za\s+([A-Z]{1,3}[a-z]?)\b'
    ] for m in re.findall(p, content, re.IGNORECASE)]
    return [r for r in set(referenced) if r in CLEN_DATA_MAP]

def build_priloga1_text(namenska_raba: str) -> str:
    """Sestavi besedilo zahteve za Prilogo 1 za posamezno namensko rabo."""
    priloga1_data = PRILOGE.get('priloga1', {})
    if not priloga1_data: return "Priloga 1 ni na voljo."

    land_uses = priloga1_data.get('land_uses', [])
    objects = priloga1_data.get('objects', [])
    
    try:
        raba_index = -1
        for i, use in enumerate(land_uses):
            if namenska_raba.upper() in use.upper().replace(" ", ""):
                raba_index = i
                break
        if raba_index == -1: return f"Namenska raba '{namenska_raba}' ni najdena v Prilogi 1."
    except ValueError:
        return f"Namenska raba '{namenska_raba}' ni najdena v Prilogi 1."

    lines = [f"Za namensko rabo '{namenska_raba}' so dovoljeni naslednji enostavni/nezahtevni objekti:\n"]
    referenced_nrp = set()
    all_nrp_conditions = {k: v for obj in objects for k, v in obj.get('nrp_conditions', {}).items()}

    for obj in objects:
        lines.append(f"**{obj['title']}**")
        for subtype in obj.get('subtypes', []):
            p_char = subtype['permissions'][raba_index]
            if p_char == "●": p_text = "Dovoljeno po splošnih določilih."
            elif p_char == "x": p_text = "Ni dovoljeno."
            else:
                p_text = f"Dovoljeno pod posebnim pogojem št. {p_char}."
                referenced_nrp.add(p_char)
            
            subtype_desc = subtype['name'] if subtype['name'] else obj['description']
            lines.append(f"- *{subtype_desc}*: {p_text}")
    
    if referenced_nrp:
        lines.append("\n**Legenda navedenih posebnih pogojev (NRP):**")
        for nrp_num in sorted(referenced_nrp):
            lines.append(f"- **Pogoj {nrp_num}**: {all_nrp_conditions.get(nrp_num, 'Opis ni na voljo.')}")
    
    return "\n".join(lines)

def build_requirements_from_db(eup_list: List[str], raba_list: List[str], project_text: str) -> List[Dict[str, Any]]:
    """Dinamično sestavi seznam zahtev glede na vsebino projekta."""
    zahteve, dodani_cleni, dodane_namenske_rabe = [], set(), set()
    splosni_pogoji_katalog = OPN_KATALOG.get("splosni_prostorski_izvedbeni_pogoji", {})

    def add_podrobni_pogoji(raba_key, kategorija):
        raba_key = raba_key.upper()
        # Preveri, ali raba obstaja v CLEN_DATA_MAP
        if raba_key in dodane_namenske_rabe or raba_key not in CLEN_DATA_MAP: return
        
        clen_data = CLEN_DATA_MAP.get(raba_key)
        
        naslov = f"{clen_data['parent_clen_key'].replace('_clen', '')}. člen - {clen_data['podrocje_naziv']} ({raba_key})"
        content = format_structured_content(clen_data['content_structured'])
        zahteve.append({"kategorija": kategorija, "naslov": naslov, "besedilo": content})
        dodane_namenske_rabe.add(raba_key)
        dodani_cleni.add(clen_data['parent_clen_key'])
        print(f"   -> Dodajam podrobno zahtevo: {naslov}")

        for ref_raba in extract_referenced_namenske_rabe(content):
            if ref_raba not in raba_list: # Preveri referenco, le če ni že med osnovnimi rabami
                print(f"   -> Namenska raba '{raba_key}' se sklicuje na: {ref_raba}")
                add_podrobni_pogoji(ref_raba, kategorija + " - Napotilo")

    print("🔎 Sestavljam seznam zahtev...")

    # KORAK 1 & 2: Splošni pogoji
    print("   -> Dodajam splošne pogoje...")
    for i in range(52, 104):
        clen_key = f"{i}_clen"
        is_mandatory = i <= 66
        keyword_match, trigger_keyword = False, ""
        if not is_mandatory:
            for keyword, mapped_clen in KEYWORD_TO_CLEN.items():
                if mapped_clen == clen_key and re.search(keyword, project_text, re.IGNORECASE):
                    keyword_match, trigger_keyword = True, keyword
                    break
        
        if (is_mandatory or keyword_match) and clen_key not in dodani_cleni:
            content = splosni_pogoji_katalog.get(clen_key)
            if content:
                naslov_match = re.search(r'^\s*\(([^)]+)\)', content)
                naslov = f"{i}. člen ({naslov_match.group(1)})" if naslov_match else f"{i}. člen"
                zahteve.append({"kategorija": "Splošni prostorski izvedbeni pogoji (PIP)", "naslov": naslov, "besedilo": content})
                dodani_cleni.add(clen_key)
                reason = 'obvezen' if is_mandatory else f'ključna beseda `{trigger_keyword}`'
                print(f"      -> Dodan splošni člen: {clen_key} ({reason})")

    # KORAK 3: Podrobni pogoji (za VSE najdene/vnešene rabe)
    ciste_namenske_rabe = sorted(list(set([r.upper() for r in raba_list if r.upper() in CLEN_DATA_MAP])))
    
    for raba in ciste_namenske_rabe:
        print(f"   -> Iščem podrobne pogoje za namensko rabo '{raba}'...")
        add_podrobni_pogoji(raba, "Podrobni prostorski izvedbeni pogoji (PIP NRP)")

    # KORAK 4: Posebni pogoji (EUP) - Popolno ujemanje
    processed_eups = set()
    priloga2_entries = PRILOGE.get('priloga2', {}).get('table_entries', [])

    for eup in eup_list:
        if not eup: continue
        normalized_eup = normalize_eup(eup)
        
        if normalized_eup in processed_eups: continue

        print(f"   -> Iščem posebne pogoje za EUP '{eup}' (Popolno ujemanje)...")
        found_entry = None
        
        # Popolno ujemanje z vnosom EUP
        for entry in priloga2_entries:
            priloga_eup = normalize_eup(entry.get("enota_urejanja", ""))
            if priloga_eup == normalized_eup:
                found_entry = entry
                break
        
        if found_entry:
            pip = found_entry.get("posebni_pip", "")
            if pip and pip.strip() and pip.strip() != "—":
                eup_name = found_entry.get("enota_urejanja", "")
                zahteve.append({"kategorija": "Posebni prostorski izvedbeni pogoji (PIP EUP)","naslov": f"Posebni PIP za EUP: {eup_name}","besedilo": pip})
                print(f"   -> Dodajam POPOLNO UJEMANJE za posebne pogoje EUP: {eup_name}")
                processed_eups.add(normalized_eup)
            else:
                print(f"      -> Za EUP '{eup}' niso bili najdeni posebni pogoji v Prilogi 2.")
        else:
            print(f"      -> EUP '{eup}' ni bil najden kot POPOLNO UJEMANJE v Prilogi 2.")


    # KORAK 5 - Preverjanje Priloge 1 (Samo ENKRAT, za vse relevantne rabe)
    if ciste_namenske_rabe:
        # Zberi samo tiste rabe, za katere je v Prilogi 1 dejansko vsebina
        rabe_za_prilogo1 = [r for r in ciste_namenske_rabe if build_priloga1_text(r) != f"Namenska raba '{r}' ni najdena v Prilogi 1."]
        
        if rabe_za_prilogo1:
            print(f"   -> Sestavljam ZDRUŽENO zahtevo za Prilogo 1 za rabe: {rabe_za_prilogo1}...")
            
            # Združevanje vsebine v eno samo polje za AI
            priloga1_content = "\n\n" + "="*50 + "\n\n".join([
                f"--- Določila za {raba} --- \n{build_priloga1_text(raba)}" for raba in rabe_za_prilogo1
            ])
            
            naslov_rabe = ", ".join(rabe_za_prilogo1)
            zahteve.append({
                "kategorija": "Skladnost z Prilogo 1 (Enostavni/Nezahtevni objekti)",
                "naslov": f"Preverjanje dopustnosti enostavnih in nezahtevnih objektov za namenske rabe: {naslov_rabe}",
                "besedilo": priloga1_content
            })
        else:
            print("   -> Preskočeno: Za nobene od navedenih rab ni najdena relevantna vsebina v Prilogi 1.")

    # KORAK 6: Oštevilčenje
    for i, zahteva in enumerate(zahteve):
        zahteva["id"] = f"Z_{i}"
    print(f"✅ Sestavljenih {len(zahteve)} zahtev.")
    return zahteve

# =============================================================================
# FILE PARSING
# =============================================================================
def parse_pdf(file_bytes: bytes) -> str:
    """Prebere PDF in vrne besedilo."""
    from pypdf import PdfReader
    try:
        pdf = PdfReader(io.BytesIO(file_bytes))
        text = "".join(page.extract_text() or "" for page in pdf.pages)
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Napaka pri branju PDF: {str(e)}")

def parse_page_string(page_str: str) -> List[int]:
    """Razčleni niz, kot je '16-25, 30', v seznam številk strani (0-indeksirano)."""
    if not page_str:
        return []
    pages = set()
    parts = page_str.split(',')
    for part in parts:
        part = part.strip()
        if '-' in part:
            try:
                start, end = map(int, part.split('-'))
                if start > 0 and end >= start:
                    pages.update(range(start - 1, end))
            except ValueError:
                continue 
        else:
            try:
                page_num = int(part)
                if page_num > 0:
                    pages.add(page_num - 1)
            except ValueError:
                continue
    return sorted(list(pages))

def convert_pdf_pages_to_images(pdf_bytes: bytes, pages_to_render_str: Optional[str]) -> List[Image.Image]:
    """Pretvori izbrane strani PDF v seznam PIL slik."""
    import fitz  # PyMuPDF
    from PIL import Image
    images = []
    if not pages_to_render_str:
        return images
        
    page_numbers = parse_page_string(pages_to_render_str)
    if not page_numbers:
        return images

    print(f"🖼️ Pretvarjam strani {pages_to_render_str} v slike...")
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page_num in page_numbers:
            if 0 <= page_num < len(doc):
                page = doc.load_page(page_num)
                pix = page.get_pixmap(dpi=200) 
                img_bytes = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_bytes))
                images.append(img)
        doc.close()
        print(f"✅ Uspešno pretvorjenih {len(images)} strani.")
    except Exception as e:
        print(f"⚠️ Napaka pri pretvorbi PDF v slike: {e}")

    return images


# =============================================================================
# PROMPT
# =============================================================================
def build_prompt(project_text: str, zahteve: List[Dict[str, Any]], izrazi_text: str, uredba_text: str) -> str:
    """Zgenerira glavni prompt za AI analizo."""
    
    zahteve_text = "".join(f"\nID: {z['id']}\nZahteva: {z['naslov']}\nBesedilo zahteve: {z['besedilo']}\n---" for z in zahteve)

    return f"""
Ti si strokovnjak za preverjanje skladnosti projektne dokumentacije z občinskimi prostorskimi akti (OPN/OPPN/PIP). Tvoja naloga je, da natančno in sistematično preveriš skladnost priloženega projekta s prostorskim aktom.

**NALOGA:**
Za vsako od spodnjih zahtev preveri skladnost priložene projektne dokumentacije. Delaj po naslednjem dvostopenjskem postopku:

**1. KORAK: ANALIZA BESEDILA**
Najprej poskusi odgovoriti na čim več zahtev z uporabo **samo tekstovnega dela** projektne dokumentacije. Poišči eksplicitne navedbe, kot so površine, tlorisne dimenzije in mere, faktorji, etažnost, število parkirnih mest itd.

**2. KORAK: CILJANA ANALIZA GRAFIK OZ. SLIK**
Ko končaš z analizo besedila, uporabi priložene slike oz. grafične priloge za dva namena:
    a) **Iskanje MANJKAJOČIH podatkov:** Za vse zahteve, kjer v besedilu nisi našel odgovora, natančno preglej grafike. Posebej pozoren bodi na:
        - **Odmike od parcelnih mej:** Te so skoraj vedno samo na situaciji.
        - **Višinske kote (terena, objekta, slemena):** Te so običajno prikazane na prerezih.
        - **Naklon strehe, višina kolenčnega zidu:** Prav tako na prerezih.
        - **Faktor zazidanosti (FZ) in faktor izrabe (FI):** Preveri, ali so na grafikah tabele s temi izračuni.
    b) **Preverjanje NESKLADIJ:** Če si v besedilu našel podatek (npr. "odmik od meje je 4.0 m"), preveri na grafiki (situaciji), ali je ta podatek skladen z vrisanim stanjem. Če odkriješ neskladje, to jasno navedi v obrazložitvi.

**RAZLAGA IZRAZOV (OPN):**
{izrazi_text or "Ni dodatnih izrazov."}

**UREDBA O RAZVRŠČANJU OBJEKTOV (KLJUČNE INFORMACIJE):**
{uredba_text or "Podatki niso na voljo."}
---
**ZAHTEVE:**
{zahteve_text}
---
**NAVODILA ZA ODGOVOR:**
Ko analiziraš zahteve in skladnost, pri podrobnih zahtevah (npr. členi 105, 106 itd...), v polje obrazložitev NUJNO kot prvo točko vnesi tudi tlorisne dimenzije stavbe oz. zunanje mere, višino in druge ključne značilnosti gradnje. 
*Primer dobrega povzetka:* "Na podlagi dokumentacije je razvidno, da so tlorisne dimenzije predmetne stanovanjske hiše 10,0 x 8,0 m. Vertikalni gabarit objekta je Pritličje + Nadstropje (P+N) z višino kolenčnega zidu 1,20 m. Streha je načrtovana kot simetrična dvokapnica z naklonom 40 stopinj, krita z opečno kritino v rdeči barvi. Fasada je predvidena v svetli, beli barvi." 
1.  Odgovori v obliki seznama (array) JSON objektov, brez kakršnegakoli drugega besedila ali markdown oznak (```json ... ```).
2.  Za VSAKO zahtevo ustvari en JSON objekt z naslednjimi polji:
    -   `"id"`: (string) ID zahteve (npr. "Z_0").
    -   `"obrazlozitev"`: (string) **IZJEMNO PODROBEN** opis ugotovitev, ki temelji na tvoji dvostopenjski analizi. Jasno loči, katere podatke si našel v besedilu in katere na grafiki. Če najdeš neskladje, ga poudari.
    -   `"evidence"`: (string) Natančna navedba vira: "Tehnično poročilo, stran X" ali "Grafika: Priloga C.2 - Situacija". Če si podatek potrdil iz obeh virov, navedi oba.
    -   `"skladnost"`: (string) Ena izmed trzech vrednosti: "Skladno", "Neskladno", ali "Ni relevantno".
    -   `"predlagani_ukrep"`: (string) Če je "Neskladno", opiši, kaj mora projektant storiti. Če je podatek manjkajoč, navedi, da ga je treba dodati. Če ukrep ni potreben, vrni "—".

3.  **POMEMBNO:** Če podatka ni ne v besedilu ne na grafikah, oceni kot "Neskladno" in v `predlagani_ukrep` zahtevaj dopolnitev dokumentacije.

4.  **!!! POSEBNO PRAVILO ZA SOGLASJA IN MNENJA ter ODMIKE !!!**
    Če zahteva omenja potrebo po pridobitvi soglasja (npr. soseda, mnenjedajalca), tvoja naloga NI preverjati, ali je bilo soglasje že pridobljeno. V takih primerih:
    -   V polje `"skladnost"` vedno vpiši **"Skladno"**.
    -   V polje `"predlagani_ukrep"` jasno navedi, katero soglasje je potrebno pridobiti.
    -   Pri navajanju odmikov v obrazložitev vnesi vse citirane odmike v dokumentaciji, tudi če so večji od 4m.
    
**Projektna dokumentacija (tekst):**
{project_text[:300000]}
---
**Projektna dokumentacija (grafične priloge):**
[Grafike so priložene in jih uporabi v drugem koraku analize za iskanje manjkajočih podatkov in preverjanje neskladij.]
""".strip()

# =============================================================================
# AI CALL
# =============================================================================
def call_gemini(prompt: str, images: List[Image.Image]) -> str:
    """2. AI Klic (Analitik): Izvede glavno analizo skladnosti z uporabo besedila in slik."""
    print("🤖 2. Klic AI (Analitik): Izvajam podrobno dvofazno analizo skladnosti...")
    try:
        model = genai.GenerativeModel(MODEL_NAME, generation_config=GEN_CFG)
        
        content_parts = [prompt]
        content_parts.extend(images)
        
        response = model.generate_content(content_parts)
        
        if not response.parts:
            reason = response.candidates[0].finish_reason if response.candidates else "NEZNAN"
            raise RuntimeError(f"Gemini ni vrnil veljavnega odgovora. Razlog: {reason}")
        text = "".join(part.text for part in response.parts)
        print(f"✅ AI Analitik odgovoril ({len(text)} znakov)")
        return text
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gemini napaka (Analitik): {str(e)}")

# =============================================================================
# PARSING AI ODGOVORA (POPRAVLJENO)
# =============================================================================
def parse_ai_response(response_text: str, expected_zahteve: List[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
    """Parsira JSON odgovor iz AI."""
    # POPRAVLJENO: Uporabljamo response_text namesto nedefinirane spremenljivke 'response'
    clean = re.sub(r'```(json)?', '', response_text, flags=re.IGNORECASE).strip()
    
    try:
        data = json.loads(clean)
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=500, detail=f"Neveljaven JSON iz AI: {e}\n\nOdgovor:\n{response_text[:500]}")
    
    results_map = {}
    if not isinstance(data, list):
        raise HTTPException(status_code=500, detail="AI ni vrnil seznama objektov v JSON formatu.")
    
    for item in data:
        if item.get("id"): results_map[item.get("id")] = item

    for z in expected_zahteve:
        if z["id"] not in results_map:
            results_map[z["id"]] = {"id": z["id"], "obrazlozitev": "AI ni uspel generirati odgovora.", "evidence": "—", "skladnost": "Neznano", "predlagani_ukrep": "Ročno preverjanje."}
            
    print(f"✅ Parsiral {len(results_map)} rezultatov.")
    return results_map

# =============================================================================
# DOCX GENERIRANJE
# =============================================================================
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches

def generate_word_report(zahteve: List[Dict[str, Any]], results_map: Dict[str, Dict[str, Any]], metadata: Dict[str, str], output_path: str) -> str:
    """Ustvari izboljšan Word dokument, strukturiran po kategorijah, z ležečo orientacijo in popravljeno sklepno ugotovitvijo."""
    print(f"📝 Ustvarjam nov Word dokument: {output_path}")
    doc = Document()
    
    # Nastavitev sekcije na ležečo (Landscape) orientacijo
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    
    # Zamenjava širine in višine (Landscape A4: 29.7cm x 21cm)
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    # Neobvezno: Prilagoditev margin za ležečo stran, npr. 1 inch (2.54cm) na vseh straneh
    margin = Inches(0.7)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin


    doc.add_paragraph(f"Datum analize: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_heading("Poročilo o skladnosti z občinskimi prostorskimi akti", level=1)

    # --- Priprava podatkov za sklepno ugotovitev ---
    neskladja = []
    for zahteva in zahteve:
        result = results_map.get(zahteva["id"], {})
        if result.get("skladnost") == "Neskladno":
            neskladja.append(zahteva.get('naslov', 'Neznan pogoj'))

    sklepni_status = "NESKLADNA" if neskladja else "SKLADNA"

    # --- Zapis sklepne ugotovitve v dokument (POPRAVLJENO BESEDILO) ---
    doc.add_heading("Sklepna ugotovitev", level=2)
    p = doc.add_paragraph()
    p.add_run("Gradnja po projektu '") # POPRAVLJENO BESEDILO
    p.add_run(metadata.get('ime_projekta', 'Ni podatka')).italic = True
    p.add_run(f"', s številko projekta '{metadata.get('stevilka_projekta', 'Ni podatka')}', " \
              f"datumom '{metadata.get('datum_projekta', 'Ni podatka')}' in projektantom " \
              f"'{metadata.get('projektant', 'Ni podatka')}', je glede na predloženo dokumentacijo ocenjena kot ")
    p.add_run(f"{sklepni_status}").bold = True
    p.add_run(" s prostorskim aktom.")

    if neskladja:
        p = doc.add_paragraph("Ugotovljena so bila neskladja v naslednjih točkah oziroma členih:")
        for tocka in neskladja:
            doc.add_paragraph(tocka, style='List Bullet')

    doc.add_paragraph() # Prazen odstavek za razmik

    # --- Generiranje tabele ---
    kategorije = {z.get("kategorija", "Ostalo"): [] for z in zahteve}
    for z in zahteve: kategorije[z.get("kategorija", "Ostalo")].append(z)
        
    preferred_order = [
        "Splošni prostorski izvedbeni pogoji (PIP)",
        "Podrobni prostorski izvedbeni pogoji (PIP NRP)",
        "Podrobni prostorski izvedbeni pogoji (PIP NRP) - Napotilo",
        "Posebni prostorski izvedbeni pogoji (PIP EUP)",
        "Skladnost z Prilogo 1 (Enostavni/Nezahtevni objekti)"
    ]
    final_order = [cat for cat in preferred_order if cat in kategorije]
    final_order.extend([cat for cat in kategorije if cat not in final_order])
    
    print(f"   -> V poročilo zapisujem naslednje kategorije: {final_order}")

    for kategorija in final_order:
        doc.add_heading(kategorija, level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text = 'Pogoj', 'Ugotovitve, obrazložitev in dokazila', 'Skladnost in ukrepi'
        for cell in hdr_cells: cell.paragraphs[0].runs[0].font.bold = True

        for zahteva in kategorije[kategorija]:
            row_cells = table.add_row().cells
            result = results_map.get(zahteva["id"], {})
            
            pogoj_p = row_cells[0].paragraphs[0]
            pogoj_p.add_run(zahteva.get('naslov', 'Brez naslova')).bold = True
            pogoj_p.add_run(f"\n\n{zahteva.get('besedilo', 'Brez besedila')}")

            obrazlozitev_p = row_cells[1].paragraphs[0]
            obrazlozitev_p.add_run("Obrazložitev:\n").bold = True
            obrazlozitev_p.add_run(result.get('obrazlozitev', '—'))
            obrazlozitev_p.add_run("\n\nDokazilo v dokumentaciji:\n").bold = True
            obrazlozitev_p.add_run(result.get('evidence', '—'))

            skladnost_p = row_cells[2].paragraphs[0]
            skladnost_p.add_run("Skladnost:\n").bold = True
            skladnost_p.add_run(result.get('skladnost', 'Neznano'))
            ukrep_text = result.get('predlagani_ukrep', '—')
            if ukrep_text and ukrep_text != "—":
                skladnost_p.add_run("\n\nPredlagani ukrepi:\n").bold = True
                skladnost_p.add_run(ukrep_text)
            
    doc.save(output_path)
    print(f"✅ Poročilo shranjeno: {output_path}")
    return output_path

# =============================================================================
# FASTAPI ENDPOINTI (POPRAVLJENO)
# =============================================================================
app = FastAPI(title="Avtomatski API za Skladnost", version="20.0.0") 
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

# Shranjevanje začasnih podatkov za uporabo v drugem koraku analize
TEMP_STORAGE: Dict[str, Dict[str, Any]] = {}
LAST_DOCX_PATH = None

@app.get("/", response_class=HTMLResponse)
def frontend():
    # POSODOBLJEN HTML IN JS za 2-stopenjski proces Z UREDLJIVIMI POLJI
    html = """
<!DOCTYPE html>
<html lang="sl">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Avtomatsko Preverjanje Skladnosti</title>
    <style>
        :root {
            --bg-gradient: linear-gradient(135deg, #eef2ff 0%, #f7f9ff 45%, #ffffff 100%);
            --card-bg: #ffffff;
            --border-color: #e3e8f1;
            --primary: #2563eb;
            --primary-dark: #1d4ed8;
            --success: #22a06b;
            --danger: #dc3545;
            --text-color: #1f2937;
            --muted: #6b7280;
            --shadow: 0 20px 40px rgba(15, 23, 42, 0.12);
            --radius-lg: 18px;
            --radius-md: 12px;
        }

        * { box-sizing: border-box; }

        body {
            margin: 0;
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background: var(--bg-gradient);
            color: var(--text-color);
            min-height: 100vh;
            display: flex;
            justify-content: center;
            align-items: stretch;
            padding: 40px 16px;
        }

        .page {
            width: min(1040px, 100%);
            display: flex;
            flex-direction: column;
            gap: 28px;
        }

        .hero {
            background: radial-gradient(circle at top left, rgba(37, 99, 235, 0.16), transparent 55%),
                        radial-gradient(circle at bottom right, rgba(34, 160, 107, 0.14), transparent 60%),
                        var(--card-bg);
            border-radius: var(--radius-lg);
            padding: 32px 36px;
            box-shadow: var(--shadow);
            display: grid;
            gap: 20px;
        }

        .hero h1 {
            font-size: clamp(2.1rem, 2vw + 1.5rem, 2.6rem);
            margin: 0;
            color: #0f172a;
        }

        .hero p {
            margin: 0;
            color: var(--muted);
            line-height: 1.6;
            font-size: 1.05rem;
        }

        .hero-steps {
            display: grid;
            gap: 14px;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
        }

        .step-card {
            background: rgba(255, 255, 255, 0.82);
            border-radius: var(--radius-md);
            padding: 16px 18px;
            border: 1px solid rgba(37, 99, 235, 0.08);
            backdrop-filter: blur(6px);
        }

        .step-card strong {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            width: 32px;
            height: 32px;
            border-radius: 50%;
            background: rgba(37, 99, 235, 0.12);
            color: var(--primary);
            margin-bottom: 10px;
            font-weight: 600;
        }

        .step-card span {
            display: block;
            color: var(--text-color);
            font-weight: 600;
            margin-bottom: 4px;
        }

        .step-card p {
            margin: 0;
            color: var(--muted);
            font-size: 0.92rem;
            line-height: 1.5;
        }

        .main-card {
            background: var(--card-bg);
            border-radius: var(--radius-lg);
            padding: 32px 36px;
            box-shadow: var(--shadow);
            display: flex;
            flex-direction: column;
            gap: 28px;
        }

        .section-header {
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 16px;
            flex-wrap: wrap;
        }

        .section-header .section-title {
            margin: 0;
        }

        .section-title {
            font-size: 1.15rem;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            color: #0f172a;
            display: flex;
            align-items: center;
            gap: 12px;
        }

        .section-title .step-badge {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            width: 34px;
            height: 34px;
            border-radius: 10px;
            background: rgba(37, 99, 235, 0.12);
            color: var(--primary);
            font-weight: 600;
        }

        .btn-inline {
            width: auto;
            padding: 12px 20px;
        }

        .upload-section {
            border: 2px dashed rgba(37, 99, 235, 0.25);
            border-radius: var(--radius-md);
            padding: 22px 24px;
            transition: all 0.25s ease;
            background: rgba(248, 250, 255, 0.8);
            display: grid;
            gap: 18px;
        }

        .upload-intro {
            display: flex;
            flex-direction: column;
            gap: 6px;
        }

        .selected-files {
            display: flex;
            flex-direction: column;
            gap: 14px;
        }

        .selected-files .file-item {
            background: rgba(255, 255, 255, 0.9);
            border: 1px solid rgba(15, 23, 42, 0.08);
            border-radius: var(--radius-md);
            padding: 14px 16px;
            box-shadow: 0 12px 24px rgba(15, 23, 42, 0.06);
        }

        .selected-files .file-head {
            display: flex;
            justify-content: space-between;
            align-items: baseline;
            gap: 12px;
            margin-bottom: 10px;
        }

        .selected-files .file-name {
            font-weight: 600;
            color: var(--text-color);
            word-break: break-word;
        }

        .selected-files .file-meta {
            color: var(--muted);
            font-size: 0.85rem;
        }

        .selected-files .file-pages label {
            font-weight: 500;
            font-size: 0.9rem;
            color: var(--muted);
        }

        .selected-files .file-pages input {
            margin-top: 6px;
        }

        .upload-section:hover,
        .upload-section.drag-active {
            border-color: var(--primary);
            background: rgba(248, 250, 255, 1);
            box-shadow: 0 16px 30px rgba(37, 99, 235, 0.12);
        }

        label {
            font-weight: 600;
            color: var(--text-color);
            display: block;
        }

        input[type=file],
        input[type=text],
        textarea {
            width: 100%;
            padding: 12px 14px;
            margin-top: 8px;
            border: 1px solid var(--border-color);
            border-radius: var(--radius-md);
            background: #f9fbff;
            font-size: 0.98rem;
            transition: border 0.2s ease, box-shadow 0.2s ease;
        }

        input[type=file] {
            padding: 12px;
            background: #fff;
        }

        input[type=text]:focus,
        textarea:focus {
            border-color: rgba(37, 99, 235, 0.65);
            box-shadow: 0 0 0 4px rgba(37, 99, 235, 0.12);
            outline: none;
        }

        textarea {
            resize: vertical;
            min-height: 70px;
            line-height: 1.5;
        }

        .subtitle {
            color: var(--muted);
            margin: 4px 0 0 0;
            line-height: 1.55;
        }

        .subtitle.muted {
            color: rgba(15, 23, 42, 0.55);
        }

        .btn {
            width: 100%;
            padding: 16px;
            border-radius: var(--radius-md);
            border: none;
            font-size: 1.05rem;
            font-weight: 600;
            color: #fff;
            background: var(--primary);
            cursor: pointer;
            transition: transform 0.2s ease, box-shadow 0.2s ease, background 0.2s ease;
        }

        .btn:hover {
            background: var(--primary-dark);
            transform: translateY(-1px);
            box-shadow: 0 12px 24px rgba(37, 99, 235, 0.2);
        }

        .btn:disabled {
            background: #a6b2d7;
            cursor: not-allowed;
            transform: none;
            box-shadow: none;
        }

        .btn-analyze {
            background: var(--success);
        }

        .btn-analyze:hover {
            background: #1c8659;
            box-shadow: 0 12px 24px rgba(34, 160, 107, 0.26);
        }

        #status {
            margin-top: 6px;
            padding: 18px 20px;
            border-radius: var(--radius-md);
            display: none;
            font-weight: 500;
            line-height: 1.5;
            border: 1px solid transparent;
        }

        #status.status-success {
            background: #ecfdf5;
            border-color: rgba(34, 160, 107, 0.3);
            color: #047857;
        }

        #status.status-error {
            background: #fef2f2;
            border-color: rgba(220, 53, 69, 0.35);
            color: #b91c1c;
        }

        #status.status-loading {
            background: #eef4ff;
            border-color: rgba(37, 99, 235, 0.35);
            color: var(--primary);
        }

        .input-group {
            border: 1px solid rgba(15, 23, 42, 0.08);
            border-radius: var(--radius-md);
            padding: 22px 24px;
            background: rgba(248, 250, 255, 0.65);
            display: flex;
            flex-direction: column;
            gap: 18px;
        }

        .manual-input-pair {
            display: grid;
            gap: 12px;
            grid-template-columns: repeat(auto-fit, minmax(220px, 1fr));
            align-items: center;
        }

        .manual-input-pair button {
            justify-self: start;
        }

        .add-btn {
            padding: 8px 14px;
            background: rgba(37, 99, 235, 0.1);
            color: var(--primary);
            border: none;
            border-radius: var(--radius-md);
            font-weight: 600;
            cursor: pointer;
            transition: all 0.2s ease;
        }

        .add-btn:hover {
            background: rgba(37, 99, 235, 0.16);
        }

        .remove-btn {
            padding: 8px 12px;
            background: rgba(220, 53, 69, 0.12);
            color: var(--danger);
            border: none;
            border-radius: var(--radius-md);
            font-weight: 600;
            cursor: pointer;
            transition: all 0.2s ease;
        }

        .remove-btn:hover {
            background: rgba(220, 53, 69, 0.18);
        }

        .key-summary {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(220px, 1fr));
            gap: 14px;
            margin-bottom: 12px;
        }

        .key-summary .summary-item {
            background: rgba(15, 23, 42, 0.04);
            border-radius: var(--radius-md);
            padding: 14px 16px;
            border: 1px solid rgba(15, 23, 42, 0.06);
        }

        .key-summary .summary-item span {
            display: block;
            font-size: 0.8rem;
            letter-spacing: 0.04em;
            text-transform: uppercase;
            color: rgba(15, 23, 42, 0.65);
        }

        .key-summary .summary-item strong {
            display: block;
            margin-top: 6px;
            font-size: 1rem;
            color: #0f172a;
            font-weight: 700;
            word-break: break-word;
        }

        .key-summary .summary-item.empty strong {
            color: rgba(15, 23, 42, 0.55);
            font-style: italic;
            font-weight: 500;
        }

        .key-data-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(240px, 1fr));
            gap: 18px;
        }

        .key-data-grid .data-item label {
            font-weight: 600;
            font-size: 0.95rem;
            margin-bottom: 6px;
        }

        .results-section {
            margin-top: 28px;
            border: 1px solid rgba(15, 23, 42, 0.08);
            border-radius: var(--radius-lg);
            background: rgba(255, 255, 255, 0.95);
            padding: 24px 26px;
            box-shadow: 0 16px 32px rgba(15, 23, 42, 0.12);
        }

        .results-section + .revision-upload {
            margin-top: 18px;
        }

        .results-header h3 {
            margin: 0 0 6px 0;
            font-size: 1.35rem;
            color: #0f172a;
        }

        .results-table {
            width: 100%;
            border-collapse: collapse;
            margin-top: 18px;
        }

        .results-table thead {
            background: rgba(37, 99, 235, 0.08);
        }

        .results-table th,
        .results-table td {
            padding: 14px 16px;
            border-bottom: 1px solid rgba(15, 23, 42, 0.08);
            vertical-align: top;
            text-align: left;
        }

        .results-table tbody tr:hover {
            background: rgba(37, 99, 235, 0.06);
        }

        .results-table tbody tr.status-neskladno {
            background: rgba(220, 53, 69, 0.08);
        }

        .results-table tbody tr.status-neskladno:hover {
            background: rgba(220, 53, 69, 0.12);
        }

        .results-table td strong {
            display: block;
            margin-bottom: 6px;
        }

        .results-table td .category-tag {
            display: inline-flex;
            align-items: center;
            padding: 4px 8px;
            border-radius: 999px;
            background: rgba(15, 23, 42, 0.08);
            font-size: 0.78rem;
            text-transform: uppercase;
            letter-spacing: 0.05em;
            color: rgba(15, 23, 42, 0.6);
        }

        .status-pill {
            display: inline-flex;
            align-items: center;
            gap: 6px;
            padding: 6px 12px;
            border-radius: 999px;
            font-weight: 600;
            font-size: 0.9rem;
        }

        .status-pill::before {
            content: '';
            display: inline-block;
            width: 8px;
            height: 8px;
            border-radius: 50%;
        }

        .status-pill.skladno {
            background: rgba(34, 160, 107, 0.15);
            color: var(--success);
        }

        .status-pill.skladno::before { background: var(--success); }

        .status-pill.neskladno {
            background: rgba(220, 53, 69, 0.18);
            color: var(--danger);
        }

        .status-pill.neskladno::before { background: var(--danger); }

        .status-pill.ni-relevantno {
            background: rgba(37, 99, 235, 0.14);
            color: var(--primary);
        }

        .status-pill.ni-relevantno::before { background: var(--primary); }

        .status-pill.neznano {
            background: rgba(15, 23, 42, 0.12);
            color: rgba(15, 23, 42, 0.8);
        }

        .status-pill.neznano::before { background: rgba(15, 23, 42, 0.7); }

        .result-note {
            color: rgba(15, 23, 42, 0.72);
            font-size: 0.92rem;
            line-height: 1.45;
        }

        .results-actions {
            margin-top: 20px;
            display: flex;
            flex-wrap: wrap;
            gap: 12px;
            justify-content: flex-start;
        }

        .btn-tertiary {
            background: rgba(37, 99, 235, 0.08);
            color: var(--primary);
        }

        .btn-tertiary:hover {
            background: rgba(37, 99, 235, 0.14);
        }

        .info-banner {
            display: none;
            justify-content: space-between;
            align-items: center;
            gap: 16px;
            padding: 18px 22px;
            border-radius: var(--radius-md);
            border: 1px solid rgba(37, 99, 235, 0.18);
            background: rgba(37, 99, 235, 0.1);
            color: #0f172a;
        }

        .info-banner .banner-text {
            display: flex;
            flex-direction: column;
            gap: 4px;
        }

        .info-banner .banner-text span {
            color: var(--muted);
            font-size: 0.9rem;
        }

        .info-banner .banner-actions {
            display: flex;
            gap: 10px;
            flex-shrink: 0;
        }

        .revision-upload {
            border: 1px dashed rgba(37, 99, 235, 0.3);
            border-radius: var(--radius-lg);
            padding: 24px 26px;
            background: rgba(37, 99, 235, 0.05);
            display: none;
            flex-direction: column;
            gap: 16px;
        }

        .revision-upload h4 {
            margin: 0;
            font-size: 1.15rem;
            color: #0f172a;
        }

        .revision-upload .revision-fields {
            display: grid;
            gap: 8px;
        }

        .revision-upload label {
            font-weight: 600;
            color: #0f172a;
        }

        .revision-upload input[type="file"],
        .revision-upload input[type="text"] {
            padding: 10px 12px;
            border: 1px solid rgba(15, 23, 42, 0.12);
            border-radius: var(--radius-md);
            font-size: 1rem;
            background: #fff;
        }

        .revision-actions {
            display: flex;
            align-items: center;
            gap: 12px;
            flex-wrap: wrap;
        }

        .revision-note {
            font-size: 0.9rem;
            color: var(--muted);
        }

        .btn-secondary {
            background: rgba(15, 23, 42, 0.08);
            color: var(--text-color);
        }

        .btn-secondary:hover {
            background: rgba(15, 23, 42, 0.14);
            color: var(--text-color);
        }

        footer {
            text-align: center;
            color: var(--muted);
            font-size: 0.9rem;
            padding-bottom: 12px;
        }

        @media (max-width: 720px) {
            body {
                padding: 24px 12px 32px;
            }

            .hero,
            .main-card {
                padding: 24px;
            }

            .manual-input-pair {
                grid-template-columns: 1fr;
            }
        }
    </style>
</head>
<body>
    <div class="page">
        <header class="hero">
            <div class="hero-content">
                <h1>Avtomatsko preverjanje skladnosti</h1>
                <p>Digitalizirajte preverjanje projektne dokumentacije. Sistem analizira naložene PDF dokumente, prepozna EUP in
                   namenske rabe ter pripravi ključne gabaritne podatke za končni pregled.</p>
            </div>
            <div class="hero-steps">
                <div class="step-card">
                    <strong>1</strong>
                    <span>Naloži dokumente</span>
                    <p>Dodajte projektno dokumentacijo in, če je potrebno, označite strani z grafikami.</p>
                </div>
                <div class="step-card">
                    <strong>2</strong>
                    <span>Preglej osnutek</span>
                    <p>AI predlog EUP/rab in ključnih podatkov je pripravljen za ročni popravek.</p>
                </div>
                <div class="step-card">
                    <strong>3</strong>
                    <span>Potrdi in analiziraj</span>
                    <p>Potrjeni podatki se uporabijo za podrobno poročilo skladnosti.</p>
                </div>
            </div>
        </header>

        <div id="restoreBanner" class="info-banner">
            <div class="banner-text">
                <strong>Na voljo je shranjena analiza za nadaljnjo obdelavo.</strong>
                <span id="restoreTimestamp"></span>
            </div>
            <div class="banner-actions">
                <button type="button" class="btn btn-analyze" id="restoreSessionBtn">Nadaljuj z dopolnitvijo</button>
                <button type="button" class="btn btn-secondary" id="discardSessionBtn">Izbriši shranjeno</button>
            </div>
        </div>

        <main class="main-card">
            <div>
                <div class="section-header">
                    <div class="section-title"><span class="step-badge">1</span> Priprava dokumentacije</div>
                    <button type="button" class="btn btn-tertiary btn-inline" id="openSavedBtn">Odpri shranjeno analizo</button>
                </div>
                <form id="uploadForm">
                    <div class="upload-section" id="dropZone">
                        <div class="upload-intro">
                            <label for="pdfFiles">Izberi projektno dokumentacijo (PDF datoteke):</label>
                            <input type="file" id="pdfFiles" accept=".pdf" multiple required>
                            <p class="subtitle">Dodajte vse tekstualne in grafične priloge projekta. Spodaj lahko za vsako datoteko določite strani za grafični pregled.</p>
                        </div>
                        <div id="selectedFilesList" class="selected-files">
                            <p class="subtitle muted">Po dodajanju datotek lahko pri posamezni prilogi vnesete strani (npr. 2, 4-6), ki naj se pretvorijo v slike za vizualno analizo.</p>
                        </div>
                    </div>
                    <button type="submit" class="btn" id="submitBtn">Ekstrahiraj ključne podatke in EUP/Rabe</button>
                </form>
                <div id="status"></div>
            </div>

            <div>
                <div class="section-title"><span class="step-badge">2</span> Pregled in potrditev podatkov</div>
                <form id="analyzeForm" style="display:none;">
                    <input type="hidden" name="session_id" id="sessionId">
                    <input type="hidden" name="existing_results_json" id="existingResults">

                    <div class="input-group">
                        <div>
                            <label>A. EUP in Namenska raba</label>
                            <p class="subtitle">AI predlog lahko prilagodite. Vsaka vrstica predstavlja par EUP in pripadajočo namensko rabo.</p>
                        </div>
                        <div id="manualInputs"></div>
                        <button type="button" id="addEupRabaBtn" class="add-btn">+ Dodaj EUP/Rab (popravek)</button>
                        <p class="subtitle" style="margin-top:-8px;">Potrjene vrednosti bodo uporabljene v naslednjem koraku analize.</p>
                    </div>

                    <div class="input-group">
                        <div>
                            <label>B. Ključni gabaritni podatki</label>
                            <p class="subtitle">Preglejte in po potrebi popravite podatke, ki jih je sistem zaznal iz dokumentacije.</p>
                        </div>
                        <div class="key-data-grid" id="keyDataFields"></div>
                    </div>

                    <button type="submit" class="btn btn-analyze" id="analyzeBtn">Izvedi podrobno analizo in pripravi poročilo</button>
                </form>
                <div id="resultsSection" class="results-section" style="display:none;">
                    <div class="results-header">
                        <h3>Rezultati analize</h3>
                        <p class="subtitle">Preglejte spodnji povzetek. Izberite zahteve, ki so bile prej neskladne ali jih želite ponovno preveriti po popravkih, nato zaženite ponovno analizo samo zanje.</p>
                    </div>
                    <div id="resultsTable"></div>
                    <div class="results-actions">
                        <button type="button" class="btn btn-tertiary btn-inline" id="saveProgressBtn">Shrani napredek</button>
                        <button type="button" class="btn btn-secondary btn-inline" id="resetSelectionBtn">Počisti izbor</button>
                    </div>
                </div>
                <div id="revisionSection" class="revision-upload">
                    <h4>Dodaj popravljeno projektno dokumentacijo</h4>
                    <p class="subtitle">Ko prejmete dopolnjeno dokumentacijo od projektanta, jo naložite tukaj. Sistem ohrani prejšnjo analizo in ponovno preveri samo izbrane točke.</p>
                    <div class="revision-fields">
                        <label for="revisionFile">Popravljena dokumentacija (PDF datoteke)</label>
                        <input type="file" id="revisionFile" accept="application/pdf" multiple>
                    </div>
                    <div class="revision-fields">
                        <label for="revisionPages">Strani za podrobno analizo (neobvezno)</label>
                        <input type="text" id="revisionPages" placeholder="npr. 2, 4-6" autocomplete="off">
                    </div>
                    <div class="revision-actions">
                        <button type="button" class="btn btn-secondary btn-inline" id="uploadRevisionBtn">Naloži popravek</button>
                        <button type="button" class="btn btn-analyze btn-inline" id="rerunSelectedBtn">Ponovna presoja izbranih zahtev</button>
                        <span class="revision-note" id="revisionInfo"></span>
                    </div>
                </div>
            </div>
        </main>

        <footer>© YEAR_PLACEHOLDER Avtomatsko preverjanje skladnosti — razvojna različica</footer>
    </div>
<script>
    const uploadForm = document.getElementById("uploadForm"),
          pdfFilesInput = document.getElementById("pdfFiles"),
          selectedFilesList = document.getElementById("selectedFilesList"),
          analyzeForm = document.getElementById("analyzeForm"),
          status = document.getElementById("status"),
          submitBtn = document.getElementById("submitBtn"),
          manualInputs = document.getElementById("manualInputs"),
          addEupRabaBtn = document.getElementById("addEupRabaBtn"),
          keyDataFields = document.getElementById("keyDataFields"),
          resultsSection = document.getElementById("resultsSection"),
          resultsTable = document.getElementById("resultsTable"),
          rerunSelectedBtn = document.getElementById("rerunSelectedBtn"),
          resetSelectionBtn = document.getElementById("resetSelectionBtn"),
          existingResultsInput = document.getElementById("existingResults"),
          analyzeBtn = document.getElementById("analyzeBtn"),
          revisionSection = document.getElementById("revisionSection"),
          revisionFileInput = document.getElementById("revisionFile"),
          revisionPagesInput = document.getElementById("revisionPages"),
          uploadRevisionBtn = document.getElementById("uploadRevisionBtn"),
          revisionInfo = document.getElementById("revisionInfo"),
          saveProgressBtn = document.getElementById("saveProgressBtn"),
          openSavedBtn = document.getElementById("openSavedBtn"),
          restoreBanner = document.getElementById("restoreBanner"),
          restoreSessionBtn = document.getElementById("restoreSessionBtn"),
          discardSessionBtn = document.getElementById("discardSessionBtn"),
          restoreTimestamp = document.getElementById("restoreTimestamp"),
          sessionIdInput = document.getElementById("sessionId");

    let currentZahteve = [];
    let currentResultsMap = {};
    const STORAGE_KEY = "mnenjaSavedState";

    // KLJUČNI SLOVAR PODATKOV ZA DINAMIČNO GENERIRANJE POLJ
    const keyLabels = {
        'glavni_objekt': 'OBJEKT (opis funkcije)', 'vrsta_gradnje': 'VRSTA GRADNJE',
        'klasifikacija_cc_si': 'CC-SI klasifikacija', 'nezahtevni_objekti': 'Nezahtevni objekti v projektu',
        'enostavni_objekti': 'Enostavni objekti v projektu', 'vzdrzevalna_dela': 'Vzdrževalna dela / manjša rekonstrukcija',
        'parcela_objekta': 'Gradbena parcela (št.)', 'stevilke_parcel_ko': 'Vse parcele in k.o.',
        'velikost_parcel': 'Skupna velikost parcel', 'velikost_obstojecega_objekta': 'Velikost obstoječega objekta',
        'tlorisne_dimenzije': 'Nove tlorisne dimenzije', 'gabariti_etaznost': 'Novi gabarit/Etažnost',
        'faktor_zazidanosti_fz': 'Faktor zazidanosti (FZ)', 'faktor_izrabe_fi': 'Faktor izrabe (FI)',
        'zelene_povrsine': 'Zelene površine (FZP/m²)', 'naklon_strehe': 'Naklon strehe',
        'kritina_barva': 'Kritina/Barva', 'materiali_gradnje': 'Materiali gradnje (npr. les)',
        'smer_slemena': 'Smer slemena', 'visinske_kote': 'Višinske kote (k.p., k.s.)',
        'odmiki_parcel': 'Odmiki od parcelnih mej', 'komunalni_prikljucki': 'Komunalni priključki/Oskrba'
    };

    const summaryLabels = {
        'glavni_objekt': 'OBJEKT',
        'vrsta_gradnje': 'VRSTA GRADNJE',
        'klasifikacija_cc_si': 'CC-SI',
        'nezahtevni_objekti': 'NEZAHTEVNI OBJEKTI',
        'enostavni_objekti': 'ENOSTAVNI OBJEKTI',
        'vzdrzevalna_dela': 'VZDRŽEVALNA DELA'
    };

    const metadataKeys = ['ime_projekta', 'stevilka_projekta'];

    function loadSavedState() {
        if (typeof localStorage === 'undefined') { return null; }
        try {
            const raw = localStorage.getItem(STORAGE_KEY);
            if (!raw) { return null; }
            return JSON.parse(raw);
        } catch (err) {
            console.warn('Shranjene seje ni mogoče prebrati:', err);
            return null;
        }
    }

    function updateRestoreBanner() {
        if (!restoreBanner) { return; }
        const saved = loadSavedState();
        if (!saved || !saved.sessionId) {
            restoreBanner.style.display = 'none';
            return;
        }

        if (restoreTimestamp) {
            const ts = saved.timestamp ? new Date(saved.timestamp) : null;
            if (ts && !Number.isNaN(ts.getTime())) {
                restoreTimestamp.textContent = `Zadnja shranitev: ${ts.toLocaleString('sl-SI')}`;
            } else {
                restoreTimestamp.textContent = '';
            }
        }

        restoreBanner.style.display = 'flex';
    }

    function formatFileSize(bytes) {
        if (typeof bytes !== 'number' || Number.isNaN(bytes)) { return ''; }
        if (bytes === 0) { return '0 B'; }
        const units = ['B', 'KB', 'MB', 'GB'];
        let size = bytes;
        let unitIndex = 0;
        while (size >= 1024 && unitIndex < units.length - 1) {
            size /= 1024;
            unitIndex += 1;
        }
        const precision = unitIndex === 0 ? 0 : (size >= 10 ? 1 : 2);
        return `${size.toFixed(precision)} ${units[unitIndex]}`;
    }

    function renderSelectedFilesList() {
        if (!selectedFilesList) { return; }

        const files = pdfFilesInput && pdfFilesInput.files ? Array.from(pdfFilesInput.files) : [];
        selectedFilesList.innerHTML = '';

        if (!files.length) {
            const info = document.createElement('p');
            info.className = 'subtitle muted';
            info.textContent = 'Po dodajanju datotek lahko pri posamezni prilogi vnesete strani (npr. 2, 4-6), ki naj se pretvorijo v slike za vizualno analizo.';
            selectedFilesList.appendChild(info);
            return;
        }

        files.forEach((file, index) => {
            const item = document.createElement('div');
            item.className = 'file-item';
            item.dataset.index = String(index);
            item.dataset.filename = file.name || `Dokument_${index + 1}`;

            const head = document.createElement('div');
            head.className = 'file-head';

            const nameSpan = document.createElement('span');
            nameSpan.className = 'file-name';
            nameSpan.textContent = file.name || `Dokument ${index + 1}`;
            head.appendChild(nameSpan);

            if (typeof file.size === 'number') {
                const metaSpan = document.createElement('span');
                metaSpan.className = 'file-meta';
                metaSpan.textContent = formatFileSize(file.size);
                head.appendChild(metaSpan);
            }

            const pagesWrapper = document.createElement('div');
            pagesWrapper.className = 'file-pages';

            const label = document.createElement('label');
            const inputId = `filePages_${index}`;
            label.setAttribute('for', inputId);
            label.textContent = 'Strani za pretvorbo v slike (neobvezno)';

            const input = document.createElement('input');
            input.type = 'text';
            input.id = inputId;
            input.className = 'file-pages-input';
            input.placeholder = 'npr. 2, 4-6';
            input.dataset.filename = file.name || `Dokument_${index + 1}`;

            pagesWrapper.appendChild(label);
            pagesWrapper.appendChild(input);

            item.appendChild(head);
            item.appendChild(pagesWrapper);

            selectedFilesList.appendChild(item);
        });
    }

    function collectCurrentState() {
        if (!sessionIdInput || !sessionIdInput.value) { return null; }
        const eupInputs = manualInputs ? Array.from(manualInputs.querySelectorAll('input[name="final_eup_list"]')) : [];
        const rabaInputs = manualInputs ? Array.from(manualInputs.querySelectorAll('input[name="final_raba_list"]')) : [];

        const eupList = eupInputs.map(input => input.value || '');
        const rabaList = rabaInputs.map(input => input.value || '');

        const keyData = {};
        Object.keys(keyLabels).forEach(key => {
            const field = analyzeForm ? analyzeForm.querySelector(`[name="${key}"]`) : null;
            if (field) { keyData[key] = field.value || ''; }
        });

        const metadata = {};
        metadataKeys.forEach(key => {
            const field = analyzeForm ? analyzeForm.querySelector(`[name="${key}"]`) : null;
            if (field) { metadata[key] = field.value || ''; }
        });

        let resultsMapToStore = currentResultsMap || {};
        if (existingResultsInput && existingResultsInput.value) {
            try {
                resultsMapToStore = JSON.parse(existingResultsInput.value);
            } catch (err) {
                console.warn('Ne morem razbrati shranjenih rezultatov, uporabim trenutno stanje.', err);
            }
        }

        return {
            sessionId: sessionIdInput.value,
            timestamp: new Date().toISOString(),
            eupList,
            rabaList,
            keyData,
            metadata,
            resultsMap: resultsMapToStore,
            zahteve: currentZahteve,
            existingResults: existingResultsInput ? existingResultsInput.value : null
        };
    }

    function persistState(auto = false) {
        if (typeof localStorage === 'undefined') { return; }
        const state = collectCurrentState();
        if (!state) {
            if (!auto) { showStatus('Ni podatkov za shranjevanje (manjka ID seje).', 'error'); }
            return;
        }
        try {
            localStorage.setItem(STORAGE_KEY, JSON.stringify(state));
            if (!auto) {
                showStatus('Trenutna analiza je shranjena za kasnejšo dopolnitev.', 'success');
            }
            updateRestoreBanner();
        } catch (err) {
            console.error('Shranjevanje ni uspelo:', err);
            if (!auto) { showStatus('Napaka pri shranjevanju napredka.', 'error'); }
        }
    }

    function applySavedState(state) {
        if (!state) { return; }
        if (!sessionIdInput) { return; }

        sessionIdInput.value = state.sessionId || '';

        clearManualInputs();

        const pairCount = Math.max(state.eupList ? state.eupList.length : 0, state.rabaList ? state.rabaList.length : 0);
        if (pairCount === 0) {
            addInputPair('', '');
        } else {
            for (let i = 0; i < pairCount; i++) {
                const eup = state.eupList && state.eupList[i] ? state.eupList[i] : '';
                const raba = state.rabaList && state.rabaList[i] ? state.rabaList[i] : '';
                addInputPair(eup, raba);
            }
        }

        const combinedData = Object.assign({}, state.metadata || {}, state.keyData || {});
        renderKeyDataFields(combinedData);

        Object.entries(state.keyData || {}).forEach(([key, value]) => {
            const field = analyzeForm ? analyzeForm.querySelector(`[name="${key}"]`) : null;
            if (field) { field.value = value || ''; }
        });
        Object.entries(state.metadata || {}).forEach(([key, value]) => {
            const field = analyzeForm ? analyzeForm.querySelector(`[name="${key}"]`) : null;
            if (field) { field.value = value || ''; }
        });

        analyzeForm.style.display = 'block';

        const serializedResults = state.existingResults || (state.resultsMap ? JSON.stringify(state.resultsMap) : '');
        if (existingResultsInput) {
            existingResultsInput.value = serializedResults || '';
        }

        renderResults(state.zahteve || [], state.resultsMap || {});
        showStatus('Shrajena analiza je naložena. Po potrebi naložite popravljeno dokumentacijo in izberite zahteve za ponovni pregled.', 'success');
    }

    function discardSavedState() {
        if (typeof localStorage === 'undefined') { return; }
        try {
            localStorage.removeItem(STORAGE_KEY);
        } catch (err) {
            console.warn('Brisanje shranjenega stanja ni uspelo:', err);
        }
        updateRestoreBanner();
    }

    function escapeHtml(value) {
        if (typeof value !== 'string') { return ''; }
        const map = { '&': '&amp;', '<': '&lt;', '>': '&gt;', '"': '&quot;', "'": '&#39;' };
        return value.replace(/[&<>"']/g, ch => map[ch]);
    }

    function statusToClass(text) {
        if (!text) return 'neznano';
        const normalized = text.toLowerCase().trim();
        if (normalized.includes('neskladno')) return 'neskladno';
        if (normalized.includes('skladno')) return 'skladno';
        if (normalized.includes('ni relevantno')) return 'ni-relevantno';
        return 'neznano';
    }

    function truncateText(text, maxLength = 220) {
        if (!text) return '—';
        if (text.length <= maxLength) return text;
        return text.slice(0, maxLength).trimEnd() + '…';
    }

    function renderResults(zahteve, resultsMap) {
        currentZahteve = Array.isArray(zahteve) ? zahteve : [];
        currentResultsMap = resultsMap && typeof resultsMap === 'object' ? resultsMap : {};

        if (!currentZahteve.length) {
            if (resultsSection) {
                resultsSection.style.display = 'none';
            }
            if (resultsTable) {
                resultsTable.innerHTML = '';
            }
            if (revisionSection) {
                revisionSection.style.display = 'none';
            }
            if (revisionInfo) {
                revisionInfo.textContent = '';
            }
            return;
        }

        let tableHtml = `<table class="results-table"><thead><tr><th style="width:60px;">Izberi</th><th>Zahteva</th><th style="width:180px;">Status</th><th>Obrazložitev / ukrep</th></tr></thead><tbody>`;

        currentZahteve.forEach(z => {
            const result = currentResultsMap[z.id] || {};
            const statusText = result.skladnost || 'Neznano';
            const statusClass = statusToClass(statusText);
            const noteRaw = result.predlagani_ukrep && result.predlagani_ukrep !== '—'
                ? result.predlagani_ukrep
                : truncateText(result.obrazlozitev || '—');
            const note = truncateText(noteRaw, 280);
            const escapedTitle = escapeHtml(z.naslov);
            const escapedCategory = escapeHtml(z.kategorija);
            const escapedStatus = escapeHtml(statusText);
            const escapedNote = escapeHtml(note);
            const normalizedStatus = (statusText || '').toLowerCase();
            const normalizedNote = (result.predlagani_ukrep || '').toLowerCase();
            const shouldCheck = normalizedStatus.includes('neskladno') || normalizedNote.includes('ponovna analiza');
            const checked = shouldCheck ? 'checked' : '';
            tableHtml += `
                <tr class="status-${statusClass}">
                    <td><input type="checkbox" value="${z.id}" ${checked}></td>
                    <td>
                        <strong>${escapedTitle}</strong>
                        <span class="category-tag">${escapedCategory}</span>
                    </td>
                    <td><span class="status-pill ${statusClass}">${escapedStatus}</span></td>
                    <td><div class="result-note">${escapedNote}</div></td>
                </tr>`;
        });

        tableHtml += '</tbody></table>';

        if (resultsTable) {
            resultsTable.innerHTML = tableHtml;
        }
        if (resultsSection) {
            resultsSection.style.display = 'block';
        }
        if (revisionSection) {
            revisionSection.style.display = 'flex';
        }
    }

    function getSelectedRequirementIds() {
        if (!resultsTable) return [];
        return Array.from(resultsTable.querySelectorAll('input[type="checkbox"]:checked')).map(cb => cb.value);
    }
    
    function renderKeyDataFields(data) {
        keyDataFields.innerHTML = '';
        
        // Polji metapodatkov (samo za prikaz, ne za urejanje)
        data.ime_projekta = data.ime_projekta || 'Ni podatka';
        data.stevilka_projekta = data.stevilka_projekta || 'Ni podatka';
        
        const metadataLabels = {
             'ime_projekta': 'Ime projekta', 'stevilka_projekta': 'Številka projekta',
        };

        // Prikaz metapodatkov
        let metaHtml = '';
        for (const key in metadataLabels) {
            metaHtml += `<div class="data-item"><label>${metadataLabels[key]}:</label><input type="text" name="${key}" id="${key}" value="${data[key]}" readonly style="background:#eee;"></div>`;
        }
        
        // Vstavitev metapodatkov pred keyDataFields (polja A in B)
        const metaDiv = document.createElement('div');
        metaDiv.className = "key-data-grid";
        metaDiv.style.gridTemplateColumns = '1fr';
        metaDiv.innerHTML = metaHtml;

        const existingMetaDiv = analyzeForm.querySelector('.key-data-grid[style*="1fr"]');
        if (existingMetaDiv) {
            existingMetaDiv.remove();
        }
        const existingSummaryDiv = analyzeForm.querySelector('.key-summary');
        if (existingSummaryDiv) {
            existingSummaryDiv.remove();
        }
        keyDataFields.insertAdjacentElement('beforebegin', metaDiv);

        // Povzetek ključnih opisnih podatkov (read-only kartice)
        const summaryDiv = document.createElement('div');
        summaryDiv.className = 'key-summary';

        for (const key in summaryLabels) {
            const rawValue = typeof data[key] === 'string' ? data[key].trim() : '';
            const value = rawValue ? rawValue : 'Ni podatka v dokumentaciji';
            const item = document.createElement('div');
            item.className = 'summary-item';
            if (value === 'Ni podatka v dokumentaciji') {
                item.classList.add('empty');
            }
            const labelSpan = document.createElement('span');
            labelSpan.textContent = summaryLabels[key];
            const valueStrong = document.createElement('strong');
            valueStrong.textContent = value;
            item.appendChild(labelSpan);
            item.appendChild(valueStrong);
            summaryDiv.appendChild(item);
        }

        metaDiv.insertAdjacentElement('afterend', summaryDiv);


        // Prikaz in urejanje razširjenih ključnih podatkov
        for (const key in keyLabels) {
            const label = keyLabels[key];
            const value = data[key] || "Ni podatka v dokumentaciji";
            const isTextArea = ['stevilke_parcel_ko', 'odmiki_parcel', 'komunalni_prikljucki', 'nezahtevni_objekti', 'enostavni_objekti', 'vzdrzevalna_dela'].includes(key);

            const div = document.createElement("div");
            div.className = "data-item";

            const labelEl = document.createElement("label");
            labelEl.textContent = label;
            labelEl.setAttribute('for', key);

            if (isTextArea) {
                 const inputEl = document.createElement("textarea");
                 inputEl.name = key;
                 inputEl.id = key;
                 inputEl.value = value;
                 div.appendChild(labelEl);
                 div.appendChild(inputEl);
            } else {
                const inputEl = document.createElement("input");
                inputEl.type = "text";
                inputEl.name = key;
                inputEl.id = key;
                inputEl.value = value;
                div.appendChild(labelEl);
                div.appendChild(inputEl);
            }
            keyDataFields.appendChild(div);
        }
    }

    function addInputPair(e = "", t = "") {
        const n = document.createElement("div");
        n.className = "manual-input-pair";
        n.innerHTML = `<input type="text" name="final_eup_list" placeholder="EUP (npr. LI-08)" value="${e}"><input type="text" name="final_raba_list" placeholder="Namenska raba (npr. SSe)" value="${t}"><button type="button" class="remove-btn">X</button>`;
        n.querySelector(".remove-btn").addEventListener("click", () => { n.remove() });
        manualInputs.appendChild(n);
    }
    addEupRabaBtn.addEventListener("click", () => addInputPair());

    function clearManualInputs() {
        manualInputs.innerHTML = '';
        keyDataFields.innerHTML = ''; // Počisti tudi polja za ključne podatke
        const existingMetaDiv = analyzeForm.querySelector('.key-data-grid[style*="1fr"]');
        if (existingMetaDiv) {
            existingMetaDiv.remove();
        }
        const existingSummaryDiv = analyzeForm.querySelector('.key-summary');
        if (existingSummaryDiv) {
            existingSummaryDiv.remove();
        }
        if (resultsSection) {
            resultsSection.style.display = 'none';
        }
        if (resultsTable) {
            resultsTable.innerHTML = '';
        }
        if (revisionSection) {
            revisionSection.style.display = 'none';
        }
        if (revisionInfo) {
            revisionInfo.textContent = '';
        }
        if (revisionFileInput) {
            revisionFileInput.value = '';
        }
        if (revisionPagesInput) {
            revisionPagesInput.value = '';
        }
        existingResultsInput.value = '';
        currentZahteve = [];
        currentResultsMap = {};
    }

    async function showStatus(e, t) {
        status.innerHTML = e;
        status.className = `status-${t}`;
        status.style.display = "block";
    }
    
    // --- KORAK 1: Ekstrakcija podatkov ---
    if (pdfFilesInput) {
        pdfFilesInput.addEventListener("change", renderSelectedFilesList);
        renderSelectedFilesList();
    }

    uploadForm.addEventListener("submit", async e => {
        e.preventDefault();

        const pdfFiles = pdfFilesInput && pdfFilesInput.files ? Array.from(pdfFilesInput.files) : [];
        if (!pdfFiles.length) { showStatus("Prosim naloži vsaj eno PDF datoteko!", "error"); return; }

        const formData = new FormData();
        const filesMeta = [];

        pdfFiles.forEach((file, index) => {
            formData.append('pdf_files', file);
            let pagesValue = '';
            if (selectedFilesList) {
                const row = selectedFilesList.querySelector(`.file-item[data-index="${index}"]`);
                if (row) {
                    const input = row.querySelector('.file-pages-input');
                    if (input && typeof input.value === 'string' && input.value.trim()) {
                        pagesValue = input.value.trim();
                    }
                }
            }
            filesMeta.push({
                name: file.name || `Dokument_${index + 1}`,
                pages: pagesValue
            });
        });

        try {
            formData.append('files_meta_json', JSON.stringify(filesMeta));
        } catch (err) {
            console.warn('Ne morem pripraviti meta podatkov o datotekah:', err);
        }

        showStatus("Analiziram dokumente in ekstrahiram ključne podatke...", "loading");
        submitBtn.disabled = true;

        try {
            const response = await fetch("/extract-data", { method: "POST", body: formData });
            
            if (!response.ok) {
                const error = await response.json();
                throw new Error(error.detail || "Napaka pri ekstrahiranju podatkov.");
            }
            
            const data = await response.json();
            
            document.getElementById('sessionId').value = data.session_id;

            // 1. Prikaz AI zaznave EUP/Raba za lažji popravek
            clearManualInputs();
            const maxLength = Math.max(data.eup.length, data.namenska_raba.length);
            for (let i = 0; i < maxLength; i++) {
                const eup = data.eup[i] || '';
                const raba = data.namenska_raba[i] || '';
                addInputPair(eup, raba);
            }
            if (maxLength === 0) {
                 addInputPair('', ''); // Dodaj vsaj eno prazno polje, če AI nič ne najde
            }

            // 2. Prikaz/renderiranje urejevalnih polj za ključne podatke
            renderKeyDataFields(data);
            
            analyzeForm.style.display = 'block';
            showStatus("Ekstrakcija podatkov uspešna. Prosim, preglejte in po potrebi POPRAVITE podatke.", "success");
            
        } catch (error) {
            showStatus(`Napaka pri ekstrakciji: ${error.message}`, "error");
            analyzeForm.style.display = 'none';
        } finally {
            submitBtn.disabled = false;
        }
    });

    async function runAnalysis(extraPayload = {}, options = {}) {
        const { isRerun = false } = options;
        const sessionId = document.getElementById('sessionId').value;
        if (!sessionId) {
            showStatus("Seja ni aktivna. Prosim, ponovite Korak 1.", "error");
            return;
        }

        const formData = new FormData(analyzeForm);

        if (Array.isArray(extraPayload.selectedIds) && extraPayload.selectedIds.length) {
            formData.append('selected_ids_json', JSON.stringify(extraPayload.selectedIds));
        }

        if (existingResultsInput) {
            const existingValue = existingResultsInput.value || '';
            if (typeof formData.set === 'function') {
                formData.set('existing_results_json', existingValue);
            } else if (existingValue) {
                formData.append('existing_results_json', existingValue);
            }
        }

        showStatus(isRerun ? "Ponovno preverjam izbrane zahteve..." : "Izvajam podrobno analizo in generiram poročilo...", "loading");
        analyzeBtn.disabled = true;
        submitBtn.disabled = true;
        if (rerunSelectedBtn) { rerunSelectedBtn.disabled = true; }

        try {
            const response = await fetch("/analyze-report", { method: "POST", body: formData });

            if (!response.ok) {
                const error = await response.json();
                throw new Error(error.detail || "Napaka pri analizi in generiranju poročila.");
            }

            const result = await response.json();
            const analyzed = result.total_analyzed ?? result.total ?? 0;
            const totalAvailable = result.total_available ?? result.total ?? analyzed;
            const scopeLabel = result.analysis_scope === "partial" ? "Ponovno analiziranih" : "Analiziranih";

            showStatus(`Poročilo uspešno ustvarjeno! ${scopeLabel} ${analyzed} od ${totalAvailable} zahtev. <br><a href="/download" style="font-weight:bold;">Prenesi poročilo (.docx)</a>`, "success");

            if (result.results_map) {
                try {
                    existingResultsInput.value = JSON.stringify(result.results_map);
                } catch (err) {
                    console.warn('Ne morem serializirati rezultatov:', err);
                    existingResultsInput.value = '';
                }
            } else {
                existingResultsInput.value = '';
            }

            renderResults(result.zahteve || [], result.results_map || {});
            persistState(true);
        } catch (error) {
            showStatus(`Kritična napaka pri analizi: ${error.message}`, "error");
        } finally {
            analyzeBtn.disabled = false;
            submitBtn.disabled = false;
            if (rerunSelectedBtn) { rerunSelectedBtn.disabled = false; }
        }
    }

    // --- KORAK 2: Izvedba analize ---
    analyzeForm.addEventListener("submit", async e => {
        e.preventDefault();
        await runAnalysis();
    });

    if (rerunSelectedBtn) {
        rerunSelectedBtn.addEventListener("click", async () => {
            if (!existingResultsInput.value) {
                showStatus("Za ponovno preverjanje potrebujemo rezultate zadnje analize.", "error");
                return;
            }

            const selected = getSelectedRequirementIds();
            if (!selected.length) {
                showStatus("Izberite vsaj eno zahtevo za ponovno preverjanje.", "error");
                return;
            }

            await runAnalysis({ selectedIds: selected }, { isRerun: true });
        });
    }

    if (resetSelectionBtn) {
        resetSelectionBtn.addEventListener("click", () => {
            if (!resultsTable) return;
            resultsTable.querySelectorAll('input[type="checkbox"]').forEach(cb => { cb.checked = false; });
        });
    }

    if (saveProgressBtn) {
        saveProgressBtn.addEventListener("click", () => {
            persistState(false);
        });
    }

    if (openSavedBtn) {
        openSavedBtn.addEventListener("click", () => {
            const saved = loadSavedState();
            if (!saved) {
                showStatus('Ni shranjenih analiz za odprtje.', 'error');
                return;
            }
            applySavedState(saved);
            if (analyzeForm && typeof analyzeForm.scrollIntoView === 'function') {
                analyzeForm.scrollIntoView({ behavior: 'smooth', block: 'start' });
            }
        });
    }

    if (restoreSessionBtn) {
        restoreSessionBtn.addEventListener("click", () => {
            const saved = loadSavedState();
            if (!saved) {
                showStatus('Ni shranjene analize za obnovitev.', 'error');
                return;
            }
            applySavedState(saved);
        });
    }

    if (discardSessionBtn) {
        discardSessionBtn.addEventListener("click", () => {
            discardSavedState();
            showStatus('Shranjeni podatki so odstranjeni.', 'success');
        });
    }

    if (uploadRevisionBtn) {
        uploadRevisionBtn.addEventListener("click", async () => {
            if (!sessionIdInput || !sessionIdInput.value) {
                showStatus('Aktivna seja ni na voljo. Ponovite prvi korak.', 'error');
                return;
            }

            const files = revisionFileInput && revisionFileInput.files ? Array.from(revisionFileInput.files) : [];
            if (!files.length) {
                showStatus('Izberite popravljeno projektno dokumentacijo v PDF formatu.', 'error');
                return;
            }

            const formData = new FormData();
            formData.append('session_id', sessionIdInput.value);
            files.forEach(file => formData.append('revision_files', file));
            if (revisionPagesInput && revisionPagesInput.value.trim()) {
                formData.append('revision_pages', revisionPagesInput.value.trim());
            }

            showStatus('Nalaganje popravljenega dokumenta in osveževanje podatkov...', 'loading');
            uploadRevisionBtn.disabled = true;

            try {
                const response = await fetch('/upload-revision', { method: 'POST', body: formData });

                if (!response.ok) {
                    let detail = 'Napaka pri nalaganju popravka.';
                    try {
                        const error = await response.json();
                        detail = error.detail || detail;
                    } catch (_) { /* ignore parsing napake */ }
                    throw new Error(detail);
                }

                const result = await response.json();
                const infoMessage = result.message || 'Popravek je uspešno naložen. Izberite neskladne zahteve in zaženite ponovno analizo.';
                showStatus(infoMessage, 'success');

                if (revisionInfo) {
                    if (result.last_revision) {
                        const ts = result.last_revision.uploaded_at ? new Date(result.last_revision.uploaded_at) : null;
                        const formatted = ts && !Number.isNaN(ts.getTime()) ? ts.toLocaleString('sl-SI') : '';
                        const uploadedNames = Array.isArray(result.last_revision.filenames) ? result.last_revision.filenames : [];
                        let label = '';
                        if (uploadedNames.length === 1) {
                            label = uploadedNames[0];
                        } else if (uploadedNames.length > 1) {
                            label = `${uploadedNames.length} datotek`;
                        } else if (files.length === 1) {
                            label = files[0].name;
                        } else if (files.length > 1) {
                            label = `${files.length} datotek`;
                        } else {
                            label = 'PDF';
                        }
                        revisionInfo.textContent = `Zadnji popravek: ${label}${formatted ? ` (${formatted})` : ''}`;
                    } else {
                        revisionInfo.textContent = 'Popravljena dokumentacija je pripravljena. Izberite zahteve in ponovno zaženite analizo.';
                    }
                }

                if (revisionFileInput) { revisionFileInput.value = ''; }
                persistState(true);
            } catch (error) {
                showStatus(error.message || 'Napaka pri nalaganju popravka.', 'error');
            } finally {
                uploadRevisionBtn.disabled = false;
            }
        });
    }

    window.addEventListener("load", () => {
        updateRestoreBanner();
    });
</script>
</body></html>"""
    return html.replace("YEAR_PLACEHOLDER", str(datetime.now().year))

@app.post("/extract-data")
async def extract_data(
    pdf_files: List[UploadFile] = File(...),
    files_meta_json: Optional[str] = Form(None),
):
    """Prvi korak: Ekstrahira ključne podatke in jih shrani za kasnejšo analizo."""
    try:
        session_id = str(datetime.now().timestamp())

        print(f"\n{'='*60}\n📤 Korak 1: Ekstrakcija podatkov (ID: {session_id})\n{'='*60}\n")
        if not pdf_files:
            raise HTTPException(status_code=400, detail="Dodajte vsaj eno PDF datoteko za analizo.")

        page_overrides: Dict[str, str] = {}
        if files_meta_json:
            try:
                parsed = json.loads(files_meta_json)
                if isinstance(parsed, list):
                    for entry in parsed:
                        if isinstance(entry, dict):
                            name = entry.get("name")
                            pages = entry.get("pages")
                            if name and isinstance(pages, str) and pages.strip():
                                page_overrides[name] = pages.strip()
                else:
                    raise ValueError
            except (json.JSONDecodeError, ValueError):
                raise HTTPException(status_code=400, detail="Neveljavni podatki o straneh za grafični pregled.")

        combined_text_parts: List[str] = []
        all_images: List[Image.Image] = []
        files_manifest: List[Dict[str, Any]] = []

        for index, upload in enumerate(pdf_files):
            pdf_bytes = await upload.read()
            if not pdf_bytes:
                continue

            file_label = upload.filename or f"Dokument_{index + 1}.pdf"
            print(f"🔄 Obdelujem datoteko: {file_label} ({len(pdf_bytes)} bajtov)")

            text = parse_pdf(pdf_bytes)
            if text:
                combined_text_parts.append(f"=== VIR: {file_label} ===\n{text}")

            page_hint = page_overrides.get(upload.filename) or page_overrides.get(file_label)
            if page_hint:
                images_for_file = convert_pdf_pages_to_images(pdf_bytes, page_hint)
                if images_for_file:
                    all_images.extend(images_for_file)

            files_manifest.append({
                "filename": file_label,
                "pages": page_hint or "",
                "size": len(pdf_bytes),
            })

        if not combined_text_parts:
            raise HTTPException(status_code=400, detail="Iz naloženih datotek ni bilo mogoče prebrati besedila.")

        project_text = "\n\n".join(combined_text_parts)
        images = all_images

        # 1. AI Detektiv (EUP/Raba)
        ai_details = call_gemini_for_details(project_text, images)

        # 2. AI Arhivar (Metapodatki)
        metadata = call_gemini_for_metadata(project_text) 
        
        # 3. AI Ekstraktor (Ključni podatki) - RAZŠIRJENO
        key_data = call_gemini_for_key_data(project_text, images)
        
        # 4. Shranjevanje vseh podatkov za drugi korak
        TEMP_STORAGE[session_id] = {
            "project_text": project_text,
            "images": images,
            "metadata": metadata,
            "ai_details": ai_details,
            "key_data": key_data,
            "source_files": files_manifest,
        }

        # 5. Priprava povratnega JSON za frontend
        response_data = {
            "session_id": session_id,
            "eup": ai_details.get("eup", []),
            "namenska_raba": ai_details.get("namenska_raba", []),
            **metadata,
            **key_data, # Vključeni vsi razširjeni podatki
            "uploaded_files": files_manifest,
        }
        
        return response_data
    except Exception as e:
        import traceback
        traceback.print_exc()
        if isinstance(e, HTTPException): raise
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/upload-revision")
async def upload_revision(
    session_id: str = Form(...),
    revision_files: List[UploadFile] = File(...),
    revision_pages: Optional[str] = Form(None),
):
    """Posodobi besedilo in slike seje z novo (popravljeno) dokumentacijo."""

    if session_id not in TEMP_STORAGE:
        raise HTTPException(status_code=404, detail="Seja ni aktivna ali je potekla. Ponovite prvi korak nalaganja dokumentacije.")

    try:
        if not revision_files:
            raise HTTPException(status_code=400, detail="Dodajte popravljeno projektno dokumentacijo.")

        combined_text_parts: List[str] = []
        combined_images: List[Image.Image] = []
        revision_names: List[str] = []

        for index, revision_file in enumerate(revision_files):
            pdf_bytes = await revision_file.read()
            if not pdf_bytes:
                continue

            file_label = revision_file.filename or f"popravek_{index + 1}.pdf"
            revision_names.append(file_label)

            text = parse_pdf(pdf_bytes)
            if text:
                combined_text_parts.append(f"=== POPRAVEK: {file_label} ===\n{text}")

            if revision_pages and revision_pages.strip():
                new_images = convert_pdf_pages_to_images(pdf_bytes, revision_pages)
                if new_images:
                    combined_images.extend(new_images)

        if not combined_text_parts and not combined_images:
            raise HTTPException(status_code=400, detail="Popravljena dokumentacija ne vsebuje uporabnih podatkov za analizo.")

        data = TEMP_STORAGE.get(session_id, {})
        if not data:
            raise HTTPException(status_code=404, detail="Podatki shranjene seje niso na voljo. Ponovite Korak 1.")

        if combined_text_parts:
            data["project_text"] = "\n\n".join(combined_text_parts)
        if combined_images:
            data["images"] = combined_images
        if revision_names:
            data["source_files"] = [{"filename": name, "pages": revision_pages or ""} for name in revision_names]

        history_entry = {
            "filenames": revision_names,
            "uploaded_at": datetime.now().isoformat(),
        }
        revision_history = data.get("revision_history", [])
        revision_history.append(history_entry)
        data["revision_history"] = revision_history

        TEMP_STORAGE[session_id] = data

        return {
            "status": "success",
            "message": "Popravljena dokumentacija je naložena. Izberite neskladne zahteve in zaženite ponovno analizo.",
            "revision_count": len(revision_history),
            "last_revision": history_entry,
        }

    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/analyze-report")
async def analyze_report(
    session_id: str = Form(...),
    # Podatki A (EUP/Raba)
    final_eup_list: List[str] = Form(None),
    final_raba_list: List[str] = Form(None),
    # Podatki B (Ključni podatki) - VSE
    glavni_objekt: str = Form("Ni podatka v dokumentaciji"),
    vrsta_gradnje: str = Form("Ni podatka v dokumentaciji"),
    klasifikacija_cc_si: str = Form("Ni podatka v dokumentaciji"),
    nezahtevni_objekti: str = Form("Ni podatka v dokumentaciji"),
    enostavni_objekti: str = Form("Ni podatka v dokumentaciji"),
    vzdrzevalna_dela: str = Form("Ni podatka v dokumentaciji"),
    parcela_objekta: str = Form("Ni podatka v dokumentaciji"),
    stevilke_parcel_ko: str = Form("Ni podatka v dokumentaciji"),
    velikost_parcel: str = Form("Ni podatka v dokumentaciji"),
    velikost_obstojecega_objekta: str = Form("Ni podatka v dokumentaciji"),
    tlorisne_dimenzije: str = Form("Ni podatka v dokumentaciji"),
    gabariti_etaznost: str = Form("Ni podatka v dokumentaciji"),
    faktor_zazidanosti_fz: str = Form("Ni podatka v dokumentaciji"),
    faktor_izrabe_fi: str = Form("Ni podatka v dokumentaciji"),
    zelene_povrsine: str = Form("Ni podatka v dokumentaciji"),
    naklon_strehe: str = Form("Ni podatka v dokumentaciji"),
    kritina_barva: str = Form("Ni podatka v dokumentaciji"),
    materiali_gradnje: str = Form("Ni podatka v dokumentaciji"),
    smer_slemena: str = Form("Ni podatka v dokumentaciji"),
    visinske_kote: str = Form("Ni podatka v dokumentaciji"),
    odmiki_parcel: str = Form("Ni podatka v dokumentaciji"),
    komunalni_prikljucki: str = Form("Ni podatka v dokumentaciji"),
    selected_ids_json: Optional[str] = Form(None),
    existing_results_json: Optional[str] = Form(None)
):
    """Drugi korak: Izvede glavno analizo s potrjenimi/popravljenimi podatki."""
    global LAST_DOCX_PATH
    
    if session_id not in TEMP_STORAGE:
        raise HTTPException(status_code=404, detail="Seja je potekla ali podatki niso bili ekstrahirani. Prosim, ponovite Korak 1.")

    data = TEMP_STORAGE.get(session_id)
    if data is None:
        raise HTTPException(status_code=404, detail="Podatki seje niso na voljo. Ponovite Korak 1.")

    try:
        print(f"\n{'='*60}\n⚙️ Korak 2: Izvajam analizo (ID: {session_id})\n{'='*60}\n")

        selected_ids: List[str] = []
        if selected_ids_json:
            try:
                parsed_ids = json.loads(selected_ids_json)
                if not isinstance(parsed_ids, list):
                    raise ValueError
                selected_ids = [str(item).strip() for item in parsed_ids if str(item).strip()]
            except (json.JSONDecodeError, ValueError):
                raise HTTPException(status_code=400, detail="Neveljaven seznam izbranih zahtev za ponovno preverjanje.")

        existing_results_map: Dict[str, Dict[str, Any]] = {}
        if existing_results_json:
            try:
                parsed_results = json.loads(existing_results_json)
                if isinstance(parsed_results, dict):
                    existing_results_map = {str(k): v for k, v in parsed_results.items() if isinstance(v, dict)}
                else:
                    raise ValueError
            except (json.JSONDecodeError, ValueError):
                raise HTTPException(status_code=400, detail="Neveljaven JSON rezultatov prejšnje analize.")

        # 1. Čiščenje in konsolidacija VSEH končnih EUP in Raba
        eup_from_form = [e.strip() for e in final_eup_list if e and e.strip()] if final_eup_list is not None else []
        raba_from_form = [r.strip().upper() for r in final_raba_list if r and r.strip()] if final_raba_list is not None else []

        final_eup_list_cleaned = list(dict.fromkeys(eup_from_form))
        final_raba_list_cleaned = list(dict.fromkeys(raba_from_form))

        if not final_raba_list_cleaned:
            raise HTTPException(status_code=404, detail="Namenska raba za analizo manjka. Prosim, vnesite jo ročno.")
            
        print(f"🔥 Finalna konsolidacija za analizo: EUP={final_eup_list_cleaned}, Raba={final_raba_list_cleaned}")
        
        # 2. Sestava zahtev
        zahteve = build_requirements_from_db(final_eup_list_cleaned, final_raba_list_cleaned, data["project_text"])

        vse_id = {z["id"] for z in zahteve}
        selected_id_set = set(selected_ids)
        if selected_id_set:
            invalid_ids = selected_id_set - vse_id
            if invalid_ids:
                raise HTTPException(status_code=400, detail=f"Izbrane zahteve ne obstajajo: {', '.join(sorted(invalid_ids))}.")
            zahteve_za_analizo = [z for z in zahteve if z["id"] in selected_id_set]
        else:
            zahteve_za_analizo = list(zahteve)

        if not zahteve_za_analizo:
            raise HTTPException(status_code=400, detail="Ni izbranih zahtev za ponovno preverjanje.")

        # 3. Priprava prompta za AI analizo - Vključitev VSEH končnih (potrjenih/popravljenih) ključnih podatkov
        final_key_data = {
            "glavni_objekt": glavni_objekt, "vrsta_gradnje": vrsta_gradnje,
            "klasifikacija_cc_si": klasifikacija_cc_si, "nezahtevni_objekti": nezahtevni_objekti,
            "enostavni_objekti": enostavni_objekti, "vzdrzevalna_dela": vzdrzevalna_dela,
            "parcela_objekta": parcela_objekta, "stevilke_parcel_ko": stevilke_parcel_ko,
            "velikost_parcel": velikost_parcel, "velikost_obstojecega_objekta": velikost_obstojecega_objekta,
            "tlorisne_dimenzije": tlorisne_dimenzije, "gabariti_etaznost": gabariti_etaznost,
            "faktor_zazidanosti_fz": faktor_zazidanosti_fz, "faktor_izrabe_fi": faktor_izrabe_fi,
            "zelene_povrsine": zelene_povrsine, "naklon_strehe": naklon_strehe,
            "kritina_barva": kritina_barva, "materiali_gradnje": materiali_gradnje,
            "smer_slemena": smer_slemena, "visinske_kote": visinske_kote,
            "odmiki_parcel": odmiki_parcel, "komunalni_prikljucki": komunalni_prikljucki
        }
        
        # Za uvodni povzetek dodamo tudi metapodatke
        metadata_formatted = "\n".join([f"- {k.replace('_', ' ').capitalize()}: {v}" for k, v in data["metadata"].items()])
        key_data_formatted = "\n".join([f"- {k.replace('_', ' ').capitalize()}: {v}" for k, v in final_key_data.items()])
        
        modified_project_text = f"""
        --- METAPODATKI PROJEKTA ---
        {metadata_formatted}
        --- KLJUČNI GABARITNI IN LOKACIJSKI PODATKI PROJEKTA (Ekstrahirano in POTRJENO) ---
        {key_data_formatted}
        --- DOKUMENTACIJA (Besedilo in grafike) ---
        {data["project_text"]}
        """
        
        # Pravilen klic funkcije: Uporabimo le 3 zahtevane argumente.
        prompt = build_prompt(modified_project_text, zahteve_za_analizo, IZRAZI_TEXT, UREDBA_TEXT)
        ai_response = call_gemini(prompt, data["images"])
        results_map = parse_ai_response(ai_response, zahteve_za_analizo)

        combined_results_map = {k: v for k, v in existing_results_map.items()}
        combined_results_map.update(results_map)

        for z in zahteve:
            if z["id"] not in combined_results_map:
                combined_results_map[z["id"]] = {
                    "id": z["id"],
                    "obrazlozitev": "Zahteva ni bila analizirana v tem koraku.",
                    "evidence": "—",
                    "skladnost": "Neznano",
                    "predlagani_ukrep": "Ponovna analiza je potrebna."
                }

        # 4. Generiranje poročila
        output_path = f"./Porocilo_Skladnosti_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        LAST_DOCX_PATH = generate_word_report(zahteve, combined_results_map, data["metadata"], output_path)

        TEMP_STORAGE[session_id] = data

        analysis_scope = "partial" if selected_id_set and len(selected_id_set) < len(zahteve) else "full"
        zahteve_summary = [{"id": z["id"], "naslov": z["naslov"], "kategorija": z.get("kategorija", "Ostalo")} for z in zahteve]

        return {
            "status": "success",
            "docx_path": LAST_DOCX_PATH,
            "total_analyzed": len(zahteve_za_analizo),
            "total_available": len(zahteve),
            "analysis_scope": analysis_scope,
            "results_map": combined_results_map,
            "zahteve": zahteve_summary
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        TEMP_STORAGE[session_id] = data
        if isinstance(e, HTTPException):
            raise
        raise HTTPException(status_code=500, detail=str(e))
